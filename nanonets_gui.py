import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import requests
import json
import os
import re
from datetime import datetime
from pathlib import Path
from PIL import Image
from docx import Document
import pandas as pd

API_KEY = "09c019f3-8348-11f1-8b50-668e6851cbef"
DEFAULT_OUTPUT_DIR = r"E:\DiskCUse\HFDownloads"

class NanoNetsOCRApp:
    def __init__(self, root):
        self.root = root
        self.root.title("NanoNets OCR 表格提取工具 - HTML → MD/Word/Excel")
        self.root.geometry("900x700")
        self.root.configure(bg='black')

        # Define fonts and colors
        self.fg_color = 'white'
        self.bg_color = 'black'
        self.entry_bg = '#2a2a2a'
        self.btn_bg = '#0066cc'
        self.btn_fg = 'white'
        self.label_font = ('Arial', 20)
        self.button_font = ('Arial', 20, 'bold')
        self.entry_font = ('Arial', 20)
        self.listbox_font = ('Courier', 20)
        self.checkbox_font = ('Arial', 20)
        self.progress_font = ('Arial', 20)
        self.log_font = ('Consolas', 18)

        # Title
        title = tk.Label(root, text="NanoNets OCR 表格提取工具", font=('Arial', 24, 'bold'),
                         fg=self.fg_color, bg=self.bg_color)
        title.pack(pady=10)

        # API Key input
        frame_key = tk.Frame(root, bg=self.bg_color)
        frame_key.pack(pady=5, fill='x', padx=20)
        tk.Label(frame_key, text="API Key:", fg=self.fg_color, bg=self.bg_color,
                 font=self.label_font).pack(side='left')
        self.api_entry = tk.Entry(frame_key, width=50, bg=self.entry_bg, fg=self.fg_color,
                                  font=self.entry_font, insertbackground=self.fg_color)
        self.api_entry.pack(side='left', padx=5)
        self.api_entry.insert(0, API_KEY)

        # Output directory selection
        frame_out = tk.Frame(root, bg=self.bg_color)
        frame_out.pack(pady=5, fill='x', padx=20)
        tk.Label(frame_out, text="輸出目錄:", fg=self.fg_color, bg=self.bg_color,
                 font=self.label_font).pack(side='left')
        self.dir_entry = tk.Entry(frame_out, width=40, bg=self.entry_bg, fg=self.fg_color,
                                  font=self.entry_font)
        self.dir_entry.pack(side='left', padx=5)
        self.dir_entry.insert(0, DEFAULT_OUTPUT_DIR)
        tk.Button(frame_out, text="瀏覽", command=self.select_dir,
                  bg=self.btn_bg, fg=self.btn_fg, font=self.button_font).pack(side='left', padx=5)

        # Image selection (multiple files)
        frame_file = tk.Frame(root, bg=self.bg_color)
        frame_file.pack(pady=10, fill='x', padx=20)
        self.file_listbox = tk.Listbox(frame_file, height=6, bg=self.entry_bg, fg=self.fg_color,
                                       selectmode=tk.EXTENDED, font=self.listbox_font,
                                       exportselection=False)
        self.file_listbox.pack(side='left', fill='both', expand=True, padx=(0, 5))
        
        btn_frame = tk.Frame(frame_file, bg=self.bg_color)
        btn_frame.pack(side='right')
        tk.Button(btn_frame, text="新增圖片", command=self.select_files,
                  bg=self.btn_bg, fg=self.btn_fg, font=self.button_font, width=12).pack(pady=2)
        tk.Button(btn_frame, text="移除選取", command=self.remove_selected,
                  bg=self.btn_bg, fg=self.btn_fg, font=self.button_font, width=12).pack(pady=2)
        tk.Button(btn_frame, text="清空列表", command=self.clear_files,
                  bg=self.btn_bg, fg=self.btn_fg, font=self.button_font, width=12).pack(pady=2)

        # Options area
        frame_options = tk.Frame(root, bg=self.bg_color)
        frame_options.pack(pady=5, fill='x', padx=20)
        
        self.save_md_var = tk.BooleanVar(value=True)
        tk.Checkbutton(frame_options, text="輸出 Markdown (.md)", variable=self.save_md_var,
                       fg=self.fg_color, bg=self.bg_color, font=self.checkbox_font,
                       selectcolor=self.entry_bg, activebackground=self.bg_color,
                       activeforeground=self.fg_color).pack(side='left', padx=10)
        
        self.save_word_var = tk.BooleanVar(value=True)
        tk.Checkbutton(frame_options, text="輸出 Word (.docx)", variable=self.save_word_var,
                       fg=self.fg_color, bg=self.bg_color, font=self.checkbox_font,
                       selectcolor=self.entry_bg, activebackground=self.bg_color,
                       activeforeground=self.fg_color).pack(side='left', padx=10)
        
        self.save_excel_var = tk.BooleanVar(value=True)
        tk.Checkbutton(frame_options, text="輸出 Excel (.xlsx)", variable=self.save_excel_var,
                       fg=self.fg_color, bg=self.bg_color, font=self.checkbox_font,
                       selectcolor=self.entry_bg, activebackground=self.bg_color,
                       activeforeground=self.fg_color).pack(side='left', padx=10)

        # Start button
        self.process_btn = tk.Button(root, text="🚀 開始辨識並轉換", command=self.process_all,
                                     bg=self.btn_bg, fg=self.btn_fg, font=self.button_font,
                                     padx=20, pady=10)
        self.process_btn.pack(pady=15)

        # Progress label
        self.progress_var = tk.StringVar(value="就緒")
        self.progress_label = tk.Label(root, textvariable=self.progress_var,
                                       fg=self.fg_color, bg=self.bg_color, font=self.progress_font)
        self.progress_label.pack()

        # Log area
        self.log_area = scrolledtext.ScrolledText(root, height=18, bg=self.bg_color, fg=self.fg_color,
                                                  font=self.log_font)
        self.log_area.pack(fill='both', expand=True, padx=20, pady=10)
        self.log_area.insert(tk.END, "就緒，請填寫 API Key、選擇輸出目錄並新增圖片。\n")
        self.log_area.config(state='disabled')

        self.image_paths = []

    def log(self, msg):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, msg + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    def update_progress(self, msg):
        self.progress_var.set(msg)
        self.root.update_idletasks()

    def select_dir(self):
        path = filedialog.askdirectory(title="選擇輸出目錄")
        if path:
            self.dir_entry.delete(0, tk.END)
            self.dir_entry.insert(0, path)

    def select_files(self):
        paths = filedialog.askopenfilenames(
            title="選擇圖片檔案",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.webp *.bmp *.tiff")]
        )
        for path in paths:
            if path not in self.image_paths:
                self.image_paths.append(path)
                self.file_listbox.insert(tk.END, os.path.basename(path))
                self.log(f"已新增: {os.path.basename(path)}")

    def remove_selected(self):
        selected = self.file_listbox.curselection()
        for idx in reversed(selected):
            self.image_paths.pop(idx)
            self.file_listbox.delete(idx)

    def clear_files(self):
        self.image_paths.clear()
        self.file_listbox.delete(0, tk.END)

    def convert_to_jpeg(self, image_path):
        """Convert any non-JPEG image to .jpeg format for Nanonets compatibility."""
        path = Path(image_path)
        suffix = path.suffix.lower()
        if suffix in ('.jpeg', '.jpg'):
            return image_path
        
        jpeg_path = path.with_suffix('.jpeg')
        
        with Image.open(image_path) as img:
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            img.save(jpeg_path, 'JPEG', quality=95)
        
        self.log(f"格式轉換: {os.path.basename(image_path)} -> {os.path.basename(jpeg_path)}")
        return str(jpeg_path)

    def extract_html_content(self, json_data):
        """從 API 返回的 JSON 中提取 markdown.content (實際上是 HTML 表格)"""
        try:
            content = json_data.get("result", {}).get("markdown", {}).get("content")
            if content is None:
                content = json_data.get("markdown", {}).get("content")
            if content is None:
                if isinstance(json_data.get("result"), list) and len(json_data["result"]) > 0:
                    first = json_data["result"][0]
                    content = first.get("markdown", {}).get("content")
            if content is None:
                raise ValueError("無法從 JSON 中找到 markdown content")
            return content
        except Exception as e:
            raise Exception(f"解析 JSON 失敗: {e}\n原始響應: {json_data}")

    def html_table_to_dataframe(self, html_str):
        """從包含 HTML 的字符串中提取第一個 <table> 並轉為 DataFrame"""
        table_pattern = re.compile(r'<table.*?>.*?</table>', re.DOTALL | re.IGNORECASE)
        table_matches = table_pattern.findall(html_str)
        if not table_matches:
            raise ValueError("未找到任何表格標籤")
        
        first_table = table_matches[0]
        try:
            dfs = pd.read_html(first_table)
            if not dfs:
                raise ValueError("解析表格後未獲得數據")
            return dfs[0]
        except Exception as e:
            raise Exception(f"解析 HTML 表格失敗: {e}\n表格片段:\n{first_table[:500]}")

    def save_as_markdown(self, content, image_path, output_dir):
        """保存 HTML/Markdown 內容到 .md 文件"""
        base_name = os.path.splitext(os.path.basename(image_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        md_filename = f"{base_name}_{timestamp}.md"
        md_path = os.path.join(output_dir, md_filename)
        
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# OCR Results - {os.path.basename(image_path)}\n\n")
            f.write(f"**Processing Time**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(content)
        return md_path

    def save_as_word(self, df, image_path, output_dir):
        """保存為 Word 文檔"""
        base_name = os.path.splitext(os.path.basename(image_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        word_filename = f"{base_name}_{timestamp}.docx"
        word_path = os.path.join(output_dir, word_filename)
        
        doc = Document()
        doc.add_heading('提取的表格數據', level=1)
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'
        for j, col_name in enumerate(df.columns):
            table.cell(0, j).text = str(col_name)
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                table.cell(i + 1, j).text = str(value)
        doc.save(word_path)
        return word_path

    def save_as_excel(self, df, image_path, output_dir):
        """保存為 Excel 文件"""
        base_name = os.path.splitext(os.path.basename(image_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_filename = f"{base_name}_{timestamp}.xlsx"
        excel_path = os.path.join(output_dir, excel_filename)
        
        df.to_excel(excel_path, index=False)
        return excel_path

    def call_nanonets_api(self, api_key, image_path):
        """調用 Nanonets OCR API (參考 111.py 的方式)"""
        url = "https://extraction-api.nanonets.com/api/v1/extract/sync"
        headers = {"Authorization": f"Bearer {api_key}"}
        
        with open(image_path, 'rb') as f:
            files = {"file": (os.path.basename(image_path), f)}
            data = {"output_format": "markdown"}
            response = requests.post(url, headers=headers, files=files, data=data, timeout=120)
        
        if response.status_code != 200:
            raise Exception(f"API 返回錯誤 (HTTP {response.status_code}): {response.text}")
        return response.json()

    def process_single(self, api_key, image_path, output_dir):
        """處理單一圖片"""
        self.log(f"\n--- 處理: {os.path.basename(image_path)} ---")
        self.update_progress(f"處理中: {os.path.basename(image_path)}")
        
        # 1. 格式轉換
        processed_path = self.convert_to_jpeg(image_path)
        
        try:
            # 2. 調用 API
            self.log("正在上傳圖片並調用 Nanonets API...")
            response_json = self.call_nanonets_api(api_key, processed_path)
            self.log("API 調用成功，正在解析回應...")
            
            # 3. 儲存原始 JSON
            base_name = os.path.splitext(os.path.basename(image_path))[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            json_filename = f"{base_name}_{timestamp}_raw.json"
            json_path = os.path.join(output_dir, json_filename)
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(response_json, f, indent=2, ensure_ascii=False)
            self.log(f"原始 JSON 已保存: {json_path}")
            
            # 4. 提取 HTML 內容
            html_content = self.extract_html_content(response_json)
            self.log(f"提取到 HTML 內容，長度: {len(html_content)} 字符")
            
            results = {}
            
            # 5. 輸出 Markdown
            if self.save_md_var.get():
                md_path = self.save_as_markdown(html_content, image_path, output_dir)
                results['md'] = md_path
                self.log(f"Markdown 已保存: {md_path}")
            
            # 6. 解析表格並輸出 Word/Excel
            if self.save_word_var.get() or self.save_excel_var.get():
                self.log("正在解析 HTML 表格...")
                df = self.html_table_to_dataframe(html_content)
                self.log(f"解析成功，表格尺寸: {df.shape[0]} 行 x {df.shape[1]} 列")
                
                if self.save_word_var.get():
                    word_path = self.save_as_word(df, image_path, output_dir)
                    results['word'] = word_path
                    self.log(f"Word 已保存: {word_path}")
                
                if self.save_excel_var.get():
                    excel_path = self.save_as_excel(df, image_path, output_dir)
                    results['excel'] = excel_path
                    self.log(f"Excel 已保存: {excel_path}")
            
            return results
            
        except Exception as e:
            raise Exception(f"處理失敗: {e}")
        finally:
            # 清理臨時檔案
            if processed_path != image_path:
                try:
                    os.remove(processed_path)
                    self.log(f"已清理臨時檔案: {os.path.basename(processed_path)}")
                except:
                    pass

    def process_all(self):
        try:
            api_key = self.api_entry.get().strip()
            if not api_key:
                messagebox.showerror("錯誤", "請先輸入 API Key")
                return
            
            if not self.image_paths:
                messagebox.showerror("錯誤", "請先新增至少一張圖片")
                return
            
            output_dir = self.dir_entry.get().strip()
            if not output_dir:
                messagebox.showerror("錯誤", "請選擇輸出目錄")
                return
            
            if not os.path.exists(output_dir):
                try:
                    os.makedirs(output_dir)
                except Exception as e:
                    messagebox.showerror("錯誤", f"無法建立輸出目錄: {e}")
                    return
            
            if not (self.save_md_var.get() or self.save_word_var.get() or self.save_excel_var.get()):
                messagebox.showerror("錯誤", "請至少勾選一種輸出格式")
                return
            
            self.process_btn.config(state='disabled')
            self.log("\n" + "="*50)
            self.log("開始批次處理...")
            self.log("="*50)
            
            success_count = 0
            fail_count = 0
            
            for i, image_path in enumerate(self.image_paths):
                self.update_progress(f"處理中 ({i+1}/{len(self.image_paths)}): {os.path.basename(image_path)}")
                try:
                    results = self.process_single(api_key, image_path, output_dir)
                    success_count += 1
                    self.log(f"✅ 完成: {os.path.basename(image_path)}")
                except Exception as e:
                    fail_count += 1
                    error_msg = f"❌ 失敗: {os.path.basename(image_path)} - {str(e)}"
                    self.log(error_msg)
            
            self.update_progress(f"完成! 成功: {success_count}, 失敗: {fail_count}")
            self.process_btn.config(state='normal')
            
            summary = f"批次處理完成!\n成功: {success_count}\n失敗: {fail_count}\n輸出目錄: {output_dir}"
            self.log("\n" + "="*50)
            self.log(summary)
            self.log("="*50)
            messagebox.showinfo("完成", summary)
        except Exception as e:
            self.log(f"發生未預期錯誤: {e}")
            messagebox.showerror("錯誤", f"發生未預期錯誤: {e}")
            self.process_btn.config(state='normal')

if __name__ == "__main__":
    try:
        root = tk.Tk()
        app = NanoNetsOCRApp(root)
        root.mainloop()
    except Exception as e:
        print(f"啟動應用程式時發生錯誤: {e}")