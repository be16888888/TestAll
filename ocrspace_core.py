#!/usr/bin/env python3
"""
OCR.Space Core Module
Contains the core OCR processing logic for OCR.Space API, separated from UI.
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Tuple, Optional, Any

import requests
import pandas as pd
from PIL import Image
from docx import Document
from docx.shared import Cm
import io
import html


def load_api_key(json_path: str = "WebOcrAPI.json") -> str:
    """Load OCR.Space API key from WebOcrAPI.json.

    Args:
        json_path: Path to the JSON configuration file.

    Returns:
        The API key string.

    Raises:
        FileNotFoundError: If the config file is missing.
        ValueError: If the API key is missing or malformed.
    """
    if not os.path.exists(json_path):
        raise FileNotFoundError(f"Configuration file not found: {json_path}")

    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            configs = json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON in {json_path}: {e}")

    if not isinstance(configs, list):
        raise ValueError(f"Expected a list in {json_path}")

    target = None
    for cfg in configs:
        if cfg.get("ModelName") == "ocr.space":
            target = cfg
            break

    if target is None:
        raise ValueError("No configuration for model 'ocr.space' found in WebOcrAPI.json")

    api_key = target.get("api_key", "").strip()
    if not api_key:
        raise ValueError("API key for 'ocr.space' is empty in WebOcrAPI.json")

    return api_key


def convert_to_png(image_path: str) -> str:
    """Convert any non-PNG image to PNG format for OCR.Space compatibility.

    Args:
        image_path: Path to the input image file.

    Returns:
        Path to the PNG image (original if already PNG, otherwise converted).
    """
    path = Path(image_path)
    suffix = path.suffix.lower()
    if suffix == '.png':
        return image_path

    png_path = path.with_suffix('.png')
    with Image.open(image_path) as img:
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        img.save(png_path, 'PNG')
    return str(png_path)


def call_ocrspace_api(api_key: str, image_path: str,
                      language: str = 'cht',
                      isOverlayRequired: bool = True,
                      isTable: bool = True,
                      OCREngine: int = 3) -> str:
    """Call OCR.Space API.

    Args:
        api_key: OCR.Space API key.
        image_path: Path to the image file to process.
        language: Language code (default 'cht' for Traditional Chinese).
        isOverlayRequired: Whether to return overlay.
        isTable: Whether to enable table recognition.
        OCREngine: OCR engine version.

    Returns:
        The extracted plain text string.

    Raises:
        Exception: If the API returns an error or the request fails.
    """
    url = "https://api.ocr.space/parse/image"
    headers = {"apikey": api_key}
    with open(image_path, 'rb') as f:
        files = {"file": (os.path.basename(image_path), f)}
        # OCR.Space API expects string values "true"/"false" for boolean parameters
        data = {
            "language": language,
            "isOverlayRequired": "true" if isOverlayRequired else "false",
            "isTable": "true" if isTable else "false",
            "OCREngine": OCREngine,
        }
        response = requests.post(url, headers=headers, files=files, data=data, timeout=120)

    if response.status_code != 200:
        raise Exception(f"API request failed (HTTP {response.status_code}): {response.text}")

    result = response.json()
    if result.get("IsErroredOnProcessing") or result.get("OCRExitCode") != 1:
        error_details = ""
        if result.get("ParsedResults"):
            parsed = result["ParsedResults"][0]
            error_details = parsed.get("ErrorDetails", "") or parsed.get("ErrorMessage", "")
        raise Exception(f"OCR processing failed: {error_details or result}")

    parsed_results = result.get("ParsedResults")
    if not parsed_results or len(parsed_results) == 0:
        raise Exception("No OCR results returned")

    parsed_text = parsed_results[0].get("ParsedText", "")
    if not parsed_text:
        raise Exception("OCR result is empty")

    # Check for common OCR failure indicators that should trigger a retry
    if "[No text detected]" in parsed_text or "キュー" in parsed_text:
        raise Exception("OCR detected no text or garbled output - trigger retry")

    return parsed_text


def extract_date_and_warehouse(text: str) -> tuple[str, str]:
    """Extract date and warehouse info from OCR text.

    Args:
        text: The OCR text.

    Returns:
        Tuple (date_str, warehouse_str). Empty strings if not found.
    """
    # Warehouse pattern: (數字庫) or (數字-數字庫)
    wh_pattern = r'\((\d+(?:-\d+)?)\s*庫\)'
    wh_match = re.search(wh_pattern, text)
    warehouse = wh_match.group(0) if wh_match else ""

    # Date pattern: xxx年xx月xx日 (可能有空格)
    date_pattern = r'(\d{2,4})\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日'
    date_match = re.search(date_pattern, text)
    date = date_match.group(0) if date_match else ""
    if date:
        date = re.sub(r'\s+', '', date)  # remove spaces

    return date, warehouse


def split_text_around_first_table(text: str) -> tuple[str, str, str]:
    """Split text into parts before, the table block, and after the first table.

    Tries HTML table first, then Markdown table.

    Args:
        text: The OCR text.

    Returns:
        Tuple (before_text, table_block, after_text). If no table, table_block and after_text are empty.
    """
    # Try HTML table
    html_pattern = re.compile(r'(<table[^>]*>.*?</table>)', re.IGNORECASE | re.DOTALL)
    html_match = html_pattern.search(text)
    if html_match:
        table_block = html_match.group(1)
        before = text[:html_match.start()]
        after = text[html_match.end():]
        before_clean = _html_to_plaintext(before)
        after_clean = _html_to_plaintext(after)
        return before_clean, table_block, after_clean

    # Try Markdown table (regex-based, matches consecutive lines containing |)
    md_pattern = re.compile(r'(^|\n)((?:[^\n]*\|[^\n]*\n)+)')
    md_match = md_pattern.search(text)
    if md_match:
        table_block = md_match.group(2).strip()
        before = text[:md_match.start()]
        after = text[md_match.end():]
        before_clean = before.strip()
        after_clean = after.strip()
        return before_clean, table_block, after_clean

    return text, "", ""


def _html_to_plaintext(html_str: str) -> str:
    """Convert HTML string to plain text (basic)."""
    decoded = html.unescape(html_str)
    # Replace <br> with newline
    plain = re.sub(r'<br\s*/?>', '\n', decoded, flags=re.IGNORECASE)
    # Remove all other tags
    plain = re.sub(r'<[^>]+>', '', plain)
    lines = [line.strip() for line in plain.splitlines() if line.strip()]
    return "\n".join(lines)


def extract_table_to_dataframe_from_block(block_text: str) -> pd.DataFrame:
    """Extract DataFrame from a table block (HTML or Markdown).

    Args:
        block_text: String containing either an HTML table or a Markdown table.

    Returns:
        pandas DataFrame.

    Raises:
        ValueError: If no table can be parsed.
    """
    if block_text.strip().startswith('<table'):
        try:
            dfs = pd.read_html(io.StringIO(block_text))
            if not dfs:
                raise ValueError("HTML table parsing yielded no data")
            return dfs[0]
        except Exception as e:
            raise Exception(f"Failed to parse HTML table: {e}")
    else:
        # Assume Markdown
        return _parse_markdown_table(block_text)


def _parse_markdown_table(md_text: str) -> pd.DataFrame:
    """Parse a Markdown table string into a DataFrame (fallback)."""
    lines = md_text.splitlines()
    data_lines = []
    for line in lines:
        if re.search(r'^\s*\|\s*[-:|]+\s*\|', line):
            continue
        if '|' in line:
            stripped = line.strip('|')
            cells = [cell.strip() for cell in stripped.split('|')]
            data_lines.append(cells)
    if not data_lines:
        raise ValueError("No valid data rows")
    header = data_lines[0]
    data = data_lines[1:]
    max_cols = max(len(row) for row in data_lines)
    for row in data:
        while len(row) < max_cols:
            row.append('')
    while len(header) < max_cols:
        header.append('')
    return pd.DataFrame(data, columns=header[:max_cols])


def fix_table_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Fix common column naming issues (e.g., first column should be '品項')."""
    if df is None or df.empty:
        return df
    df.columns = [str(col).strip() for col in df.columns]
    if len(df.columns) == 8:
        if df.columns[0] == '' or '品項' not in df.columns[0]:
            new_cols = list(df.columns)
            new_cols[0] = '品項'
            df.columns = new_cols
    return df


def save_as_word(df: pd.DataFrame, filename: str, before_text: str, after_text: str) -> str:
    """Save DataFrame as a Word document with header info.

    Args:
        df: DataFrame containing the table data.
        filename: Full path to the .docx file.
        before_text: Text before the table (used for header).
        after_text: Text after the table.

    Returns:
        The filename (for consistency).
    """
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    if before_text:
        before_lines = [line.strip() for line in before_text.split('\n') if line.strip()]
        # Merge header lines (warehouse, date, filler, unit) into one line
        merged = _merge_title_info(before_lines)
        for line in merged:
            doc.add_paragraph(line)

    if df is not None and not df.empty and df.shape[1] > 0:
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'
        for j, col_name in enumerate(df.columns):
            table.cell(0, j).text = str(col_name)
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                cell_text = "" if pd.isna(value) else str(value)
                table.cell(i + 1, j).text = cell_text

    if after_text:
        for line in after_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    doc.save(filename)
    return filename


def save_as_excel(df: pd.DataFrame, image_path: str, output_dir: str) -> str:
    """Save DataFrame as an Excel file.

    Args:
        df: DataFrame containing the table data.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved Excel file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    excel_filename = f"{base_name}_{timestamp}.xlsx"
    excel_path = os.path.join(output_dir, excel_filename)
    df.to_excel(excel_path, index=False)
    return excel_path


def save_as_text(text: str, image_path: str, output_dir: str) -> str:
    """Save plain text to a .txt file.

    Args:
        text: The text to save.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved text file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    txt_filename = f"{base_name}_{timestamp}.txt"
    txt_path = os.path.join(output_dir, txt_filename)
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text)
    return txt_path


def save_as_md(text: str, image_path: str, output_dir: str) -> str:
    """Save text to a .md file.

    Args:
        text: The text to save.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved markdown file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_filename = f"{base_name}_{timestamp}.md"
    md_path = os.path.join(output_dir, md_filename)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# OCR Results - {os.path.basename(image_path)}\n\n")
        f.write(f"**Processing Time**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write(text)
    return md_path


def _merge_title_info(lines: list[str]) -> list[str]:
    """Merge header information lines (warehouse, date, filler, unit) into a single line.

    Args:
        lines: List of lines before the table.

    Returns:
        List of lines with header info merged.
    """
    if not lines:
        return lines
    new_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if '庫)' in line and (line.startswith('#') or '(' in line):
            combined = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if any(next_line.startswith(prefix) for prefix in [
                    '**日期：**', '日期：',
                    '**填單者：**', '填單者：',
                    '**單位：**', '單位：',
                    '日期：', '填單者：', '單位：'  # also match without markdown
                ]):
                    combined.append(next_line)
                    j += 1
                else:
                    break
            merged_line = '     '.join(combined)
            new_lines.append(merged_line)
            i = j
        else:
            new_lines.append(line)
            i += 1
    return new_lines


def process_single(api_key: str, image_path: str, output_dir: str,
                   save_word: bool = True, save_excel: bool = True,
                   save_text: bool = True, save_md: bool = True,
                   language: str = 'cht', isTable: bool = True) -> Dict[str, str]:
    """Process a single image through the full OCR.Space OCR pipeline.

    Args:
        api_key: OCR.Space API key.
        image_path: Path to the image file to process.
        output_dir: Directory where output files should be saved.
        save_word: Whether to save a Word document.
        save_excel: Whether to save an Excel file.
        save_text: Whether to save a plain text file.
        save_md: Whether to save a Markdown file.
        language: Language code (default 'cht' for Traditional Chinese).
        isTable: Whether to enable table recognition.

    Returns:
        Dictionary with keys 'txt', 'md', 'word', 'excel' (if saved) mapping to file paths.

    Raises:
        Exception: If any step in the processing fails.
    """
    # 1. API call with PNG retry on failure
    max_retries = 1
    attempt = 0
    final_ocr_text = ""
    png_path = None  # to track temporary PNG
    while attempt <= max_retries:
        try:
            if attempt == 0:
                # First attempt: use original image
                processed_path = image_path
            else:
                # Retry: convert to PNG
                processed_path = convert_to_png(image_path)
                png_path = processed_path  # remember to clean up
            ocr_text = call_ocrspace_api(api_key, processed_path,
                                         language=language,
                                         isTable=isTable)
            final_ocr_text = ocr_text
            break
        except Exception as e:
            if attempt < max_retries:
                # Try converting to PNG and retry
                try:
                    png_path = convert_to_png(image_path)
                except Exception:
                    raise Exception(f"OCR.Space processing failed: {e}")
            else:
                raise Exception(f"OCR.Space processing failed: {e}")
        attempt += 1

    # Clean up temporary PNG if we created one
    if png_path and os.path.exists(png_path):
        try:
            os.remove(png_path)
        except OSError:
            pass  # Ignore cleanup errors

    # 2. Extract date and warehouse for filename
    date_str, warehouse_str = extract_date_and_warehouse(final_ocr_text)
    if date_str or warehouse_str:
        base_filename = f"{date_str}{warehouse_str}"
    else:
        base_name = os.path.splitext(os.path.basename(image_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{base_name}_{timestamp}"

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    results = {}

    # 3. Save text file (with base_filename)
    if save_text:
        txt_filename = f"{base_filename}.txt"
        txt_path = os.path.join(output_dir, txt_filename)
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(final_ocr_text)
        results['txt'] = txt_path

    # 4. Save markdown file (with base_filename)
    if save_md:
        md_filename = f"{base_filename}.md"
        md_path = os.path.join(output_dir, md_filename)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# OCR Results - {os.path.basename(image_path)}\n\n")
            f.write(f"**Processing Time**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(final_ocr_text)
        results['md'] = md_path

    # 5. Extract table + PNG retry on table failure OR column-count mismatch OR "[No text detected]"
    #    (matches original behaviour: JPG→PNG retry when table parsing yields
    #     wrong column count, e.g. not 8 columns for the stock form)
    df = None
    before_text = ""
    after_text = ""
    table_ocr_text = final_ocr_text

    max_table_retries = 1
    table_attempt = 0
    while table_attempt <= max_table_retries:
        need_retry = False
        try:
            before_text, table_block, after_text = split_text_around_first_table(table_ocr_text)
            if table_block:
                df = extract_table_to_dataframe_from_block(table_block)
                df = fix_table_columns(df)
                # Original behaviour: when the table doesn't have the expected
                # 8 columns, treat it as a bad parse and retry with PNG.
                if df is not None and not df.empty and df.shape[1] != 8:
                    need_retry = True
            else:
                df = None
            
            # Also check for OCR failure indicators that should trigger PNG retry
            if "[No text detected]" in table_ocr_text or "キュー" in table_ocr_text:
                need_retry = True
                
            if not need_retry:
                break  # Success or no table found
        except Exception as e:
            need_retry = True

        if table_attempt < max_table_retries:
            # Table parsing failed or wrong column count or OCR failure — retry with PNG
            try:
                png_path = convert_to_png(image_path)
                table_ocr_text = call_ocrspace_api(api_key, png_path,
                                                   language=language,
                                                   isTable=isTable)
                try:
                    if png_path != image_path:
                        os.remove(png_path)
                except OSError:
                    pass
            except Exception:
                df = None
                break
        else:
            df = None
            break
        table_attempt += 1

    # 6. Save Word always (even without table — matches original behaviour)
    if save_word:
        word_filename = f"{base_filename}.docx"
        word_path = save_as_word(df, os.path.join(output_dir, word_filename),
                                 before_text, after_text)
        results['word'] = word_path

    # 7. Save Excel (with base_filename, only if table exists)
    if save_excel and df is not None and not df.empty:
        excel_filename = f"{base_filename}.xlsx"
        excel_path = os.path.join(output_dir, excel_filename)
        df.to_excel(excel_path, index=False)
        results['excel'] = excel_path

    return results