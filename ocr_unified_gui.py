#!/usr/bin/env python3
"""
統一 OCR.Space / Nanonets / LlamaIndex GUI
- 左側原圖 / 右側按鈕啟動 LibreOffice Writer 開啟 Word 檔（可編輯+存檔+修改格式）
- 依序嘗試 LlamaIndex -> Nanonets -> OCR.Space，一個失敗自動切下一個
- 支援修改 API Key 並寫回 WebOcrAPI.json
- 黑底白字、預設字體 20
"""

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from tkinter import ttk
import os
import json
import tempfile
import webbrowser
import subprocess
import threading
import time
import hashlib
from pathlib import Path

from ocrspace_core import process_single as ocrspace_process, load_api_key as ocrspace_load_key
from nanonets_core import process_single as nanonets_process, load_api_key as nanonets_load_key
from llamaindex_core import process_image as llamaindex_process, load_api_key as llamaindex_load_key, get_base_url as llamaindex_get_base_url
from image_utils import convert_to_jpeg


# ---------------------------
# Path constants
# ---------------------------
_PROJECT_DIR = Path(__file__).resolve().parent
_API_KEY_FILE = str(_PROJECT_DIR / "WebOcrAPI.json")
_HASH_DB_PATH = str(_PROJECT_DIR / "image_hashes.json")
# ---------------------------
FG_COLOR = 'white'
BG_COLOR = 'black'
ENTRY_BG = '#333333'
BTN_BG = '#555555'
LOG_BG = '#1a1a1a'
LOG_FG = '#00ff00'
DEFAULT_FONT = ('Arial', 20)
TITLE_FONT = ('Arial', 20, 'bold')
NORMAL_FONT = ('Arial', 14)
SMALL_FONT = ('Arial', 12)
WORD_FILENAME_FONT = ('Arial', 16)  # Word 檔名字體 16（只顯示檔名）
MONO_FONT = ('Consolas', 16)
BOX_FONT = ('Arial', 16)  # 表格上方/下方文字框字體 16（供手改日期/庫別）

OUTPUT_DIR = r"/mnt/e/DiskCUse/HFDownloads"


# ---------------------------
# Helpers
# ---------------------------
def time_str() -> str:
    return time.strftime("%H:%M:%S")


# ---------------------------
# Main App
# ---------------------------
class UnifiedOCRApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        root.title("統一 OCR 辨識工具")
        root.geometry("1600x920")
        root.configure(bg=BG_COLOR)
        root.minsize(1200, 720)

        # state first
        self.image_paths = []        # list of paths, multi-select
        self.current_image_idx = 0   # which image is shown
        self.latest_docx_path = None
        self.latest_docx_paths = {}  # image_path -> docx_path map
        self.running = False
        self.api_keys = {}
        # hash tracking
        self.image_hashes = {}       # hash -> original_path

        # UI
        self._build_top_bar()
        self._build_middle()
        self._build_bottom_log()

        # 載入雜湊資料庫 (需在 log_area 建立後，因會呼叫 self.log)
        self._load_hash_db()

        # keys
        try:
            self._load_all_keys()
        except Exception as e:
            messagebox.showerror("啟動失敗", f"WebOcrAPI.json 讀取失敗：{e}")

    # -------------------
    # Load / Save keys
    # -------------------
    def _load_all_keys(self):
        try:
            with open(_API_KEY_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            messagebox.showerror("啟動失敗", f"WebOcrAPI.json 讀取失敗：{e}")
            data = []

        for item in data if isinstance(data, list) else []:
            name = item.get("ModelName")
            key = item.get("api_key", "")
            self.api_keys[name] = key

        self.log(f"已載入 API Key：{list(self.api_keys.keys())}")

    def _save_key_to_json(self, model_name: str, new_key: str):
        if not new_key.strip():
            messagebox.showwarning("提示", "API Key 不可為空，不會更新。")
            return False
        try:
            with open(_API_KEY_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            messagebox.showerror("錯誤", f"讀取 WebOcrAPI.json 失敗：{e}")
            return False

        updated = False
        for item in data if isinstance(data, list) else []:
            if item.get("ModelName") == model_name:
                item["api_key"] = new_key.strip()
                updated = True
                break

        if not updated:
            # append new entry
            data.append({"ModelName": model_name, "api_key": new_key.strip()})

        try:
            with open(_API_KEY_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            messagebox.showerror("錯誤", f"寫入 WebOcrAPI.json 失敗：{e}")
            return False

        self.api_keys[model_name] = new_key.strip()
        return True

    # -------------------
    # Hash tracking
    # -------------------
    def _load_hash_db(self):
        """載入 image_hashes.json，建立 hash -> path 映射"""
        try:
            with open(_HASH_DB_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            # data: [{"hash": "...", "path": "..."}, ...]
            self.image_hashes = {item["hash"]: item["path"] for item in data if isinstance(item, dict)}
            self.log(f"已載入雜湊資料庫：{len(self.image_hashes)} 筆")
        except FileNotFoundError:
            self.image_hashes = {}
            self.log("雜湊資料庫不存在，將建立新檔")
        except Exception as e:
            self.image_hashes = {}
            self.log(f"雜湊資料庫讀取失敗：{e}")

    def _save_hash_db(self):
        """將 image_hashes 寫入 image_hashes.json"""
        try:
            data = [{"hash": h, "path": p} for h, p in self.image_hashes.items()]
            with open(_HASH_DB_PATH, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            self.log(f"雜湊資料庫寫入失敗：{e}")

    @staticmethod
    def _calc_file_hash(path: str) -> str:
        """計算檔案 SHA256 雜湊"""
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()

    # -------------------
    # UI build
    # -------------------
    def _build_top_bar(self):
        bar = tk.Frame(self.root, bg=BG_COLOR)
        bar.pack(side='top', fill='x', padx=20, pady=10)

        tk.Button(
            bar, text="選擇圖片（可多選）", command=self.select_files,
            bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT, padx=20, pady=8
        ).pack(side='left', padx=10)

        # file list dropdown
        tk.Label(bar, text="圖片：", fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT).pack(side='left', padx=(10, 2))
        self.file_var = tk.StringVar(value="尚未選擇")
        self.file_combo = ttk.Combobox(
            bar, textvariable=self.file_var, values=["尚未選擇"],
            state='readonly', width=35, font=SMALL_FONT
        )
        self.file_combo.pack(side='left', padx=2)
        self.file_combo.bind('<<ComboboxSelected>>', self._on_file_selected)

        tk.Button(
            bar, text="開始辨識", command=self.start_process,
            bg='#0066cc', fg=FG_COLOR, font=SMALL_FONT, padx=24, pady=8
        ).pack(side='left', padx=10)

        tk.Button(
            bar, text="修改 API Key", command=self.open_api_manager,
            bg='#885500', fg=FG_COLOR, font=SMALL_FONT, padx=20, pady=8
        ).pack(side='left', padx=10)

        tk.Label(bar, text="輸出：", fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT).pack(side='left', padx=10)
        self.out_var = tk.StringVar(value=OUTPUT_DIR)
        tk.Entry(bar, textvariable=self.out_var, width=42, bg=ENTRY_BG, fg=FG_COLOR, font=SMALL_FONT).pack(side='left', padx=5)
        tk.Button(bar, text="瀏覽", command=self.select_out_dir, bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT).pack(side='left')

    def _build_middle(self):
        mid = tk.PanedWindow(self.root, orient=tk.HORIZONTAL, bg=BG_COLOR, sashwidth=8)
        mid.pack(fill='both', expand=True, padx=20, pady=8)

        # left: image
        left = tk.Frame(mid, bg=BG_COLOR)

        # 圖片工具列：放大 / 縮小 / 重置
        img_toolbar = tk.Frame(left, bg=BG_COLOR)
        img_toolbar.pack(fill='x', padx=0, pady=(0, 2))
        tk.Label(img_toolbar, text='原圖', fg=FG_COLOR, bg=BG_COLOR, font=TITLE_FONT).pack(side='left')
        tk.Button(
            img_toolbar, text='🔍 放大', command=self._zoom_in,
            bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT, padx=12, pady=4
        ).pack(side='left', padx=(12, 4))
        tk.Button(
            img_toolbar, text='🔍 縮小', command=self._zoom_out,
            bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT, padx=12, pady=4
        ).pack(side='left', padx=4)
        tk.Button(
            img_toolbar, text='↺ 重置', command=self._zoom_reset,
            bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT, padx=12, pady=4
        ).pack(side='left', padx=4)
        self.zoom_label = tk.Label(
            img_toolbar, text='100%', fg='#aaaaaa', bg=BG_COLOR, font=SMALL_FONT
        )
        self.zoom_label.pack(side='left', padx=8)

        self.image_canvas = tk.Canvas(left, bg='#111111', highlightthickness=1, highlightbackground='#555555')
        self.image_canvas.pack(fill='both', expand=True)
        self.image_canvas.bind('<Configure>', lambda e: self._draw_image_preview())
        self.preview_photo = None
        self._zoom_scale = 1.0
        self._img_offset_x = 0
        self._img_offset_y = 0
        self._drag_start_x = None
        self._drag_start_y = None
        self.image_canvas.bind('<ButtonPress-1>', self._on_img_drag_start)
        self.image_canvas.bind('<B1-Motion>', self._on_img_drag_move)
        self.image_canvas.bind('<ButtonRelease-1>', self._on_img_drag_end)
        # 滾輪縮放
        self.image_canvas.bind('<MouseWheel>', self._on_mousewheel)   # Windows
        self.image_canvas.bind('<Button-4>', self._on_mousewheel)     # Linux scroll up
        self.image_canvas.bind('<Button-5>', self._on_mousewheel)     # Linux scroll down
        mid.add(left, minsize=400)

        # right: Word 表格編輯區
        right = tk.Frame(mid, bg=BG_COLOR)
        # --- Toolbar ---
        toolbar = tk.Frame(right, bg=BG_COLOR)
        toolbar.pack(fill='x', padx=0, pady=(0, 4))
        tk.Label(toolbar, text="Word 表格編輯", fg=FG_COLOR, bg=BG_COLOR, font=TITLE_FONT).pack(side='left')
        self.word_status = tk.Label(
            toolbar, text="（尚未產出）", fg='orange', bg=BG_COLOR, font=SMALL_FONT
        )
        self.word_status.pack(side='left', padx=12)

        # file path display
        self.word_path_var = tk.StringVar(value="")
        path_bar = tk.Frame(right, bg=BG_COLOR)
        path_bar.pack(fill='x', pady=(0, 4))
        tk.Label(
            path_bar, textvariable=self.word_path_var,
            fg='#aaaaaa', bg=BG_COLOR, font=WORD_FILENAME_FONT, anchor='w'
        ).pack(side='left', fill='x', expand=True)

        # 表格上方文字區域 (Phase 10: 表格上方，可編輯，供日期/庫別辨識錯誤時修改)
        self.before_table_frame = tk.Frame(right, bg=BG_COLOR)
        self.before_table_label = tk.Label(
            self.before_table_frame, text="表格上方文字：", fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT
        )
        self.before_table_label.pack(anchor='w', pady=(6, 2))
        self.before_table_text = scrolledtext.ScrolledText(
            self.before_table_frame, height=4, bg=LOG_BG, fg=LOG_FG,
            font=BOX_FONT, state='normal', wrap='word'
        )
        self.before_table_text.pack(fill='x')
        self.before_table_frame.pack(fill='x', pady=(4, 0))
        def _on_before_text_modified(event):
            if self.before_table_text.edit_modified():
                self._dirty = True
                self.before_table_text.edit_modified(False)
        self.before_table_text.bind('<<Modified>>', _on_before_text_modified)

        # --- Treeview table ---
        table_frame = tk.Frame(right, bg=BG_COLOR)
        table_frame.pack(fill='both', expand=True)
        vsb = ttk.Scrollbar(table_frame, orient='vertical')
        vsb.pack(side='right', fill='y')
        # Horizontal scrollbar
        hsb = ttk.Scrollbar(table_frame, orient='horizontal')
        hsb.pack(side='bottom', fill='x')

        self.tree = ttk.Treeview(
            table_frame,
            columns=[],
            show='headings',
            yscrollcommand=vsb.set,
            xscrollcommand=hsb.set,
            selectmode='browse'
        )
        self.tree.pack(side='left', fill='both', expand=True)
        vsb.config(command=self.tree.yview)
        hsb.config(command=self.tree.xview)

        # 表格下方文字區域
        self.after_table_frame = tk.Frame(right, bg=BG_COLOR)
        self.after_table_label = tk.Label(
            self.after_table_frame, text="表格下方文字：", fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT
        )
        self.after_table_label.pack(anchor='w', pady=(6, 2))
        self.after_table_text = scrolledtext.ScrolledText(
            self.after_table_frame, height=4, bg=LOG_BG, fg=LOG_FG,
            font=BOX_FONT, state='normal', wrap='word'
        )
        self.after_table_text.pack(fill='x')
        self.after_table_frame.pack(fill='x', pady=(4, 0))
        # 下方文字被編輯時標記 dirty (含日期/下收手寫字)，確保回存時一併寫回
        def _on_after_text_modified(event):
            if self.after_table_text.edit_modified():
                self._dirty = True
                self.after_table_text.edit_modified(False)
        self.after_table_text.bind('<<Modified>>', _on_after_text_modified)

        # Style for dark theme
        style = ttk.Style()
        style.theme_use('default')
        style.configure("Treeview",
            background=LOG_BG, fieldbackground=LOG_BG, foreground=LOG_FG,
            font=MONO_FONT, rowheight=28, borderwidth=0)
        style.configure("Treeview.Heading",
            background=BTN_BG, foreground=FG_COLOR, font=SMALL_FONT, borderwidth=1)
        style.map("Treeview", background=[('selected', '#004488')])

        # Bind double-click for cell editing
        self.tree.bind('<Double-1>', self._on_cell_double_click)
        self.tree.bind('<Button-1>', self._on_tree_click)

        # --- 入庫資訊區 (Phase 3) ---
        db_row = tk.Frame(right, bg=BG_COLOR)
        db_row.pack(fill='x', pady=(8,0))
        tk.Label(db_row,text="表格日期:",fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT).grid(row=0,column=0,sticky='w',padx=(0,4))
        self.biz_date_var = tk.StringVar(value=time.strftime("%Y-%m-%d"))
        tk.Entry(db_row,textvariable=self.biz_date_var,width=12,bg=ENTRY_BG,fg=FG_COLOR,font=SMALL_FONT).grid(row=0,column=1,padx=(0,12))
        tk.Label(db_row,text="庫別:",fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT).grid(row=0,column=2,sticky='w',padx=(0,4))
        # 庫別：顯示 OCR 辨識出的文字 (如 4-1庫)，可手動編輯；入庫前自動 upsert 進 libraries 表 (FK 相容)
        self.lib_var = tk.StringVar(value="")
        tk.Entry(db_row,textvariable=self.lib_var,width=12,bg=ENTRY_BG,fg=FG_COLOR,font=SMALL_FONT).grid(row=0,column=3,padx=(0,12))
        tk.Label(db_row,text="品項:",fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT).grid(row=0,column=4,sticky='w',padx=(0,4))
        self.item_name_var = tk.StringVar()
        tk.Entry(db_row,textvariable=self.item_name_var,width=18,bg=ENTRY_BG,fg=FG_COLOR,font=SMALL_FONT).grid(row=0,column=5)
        self.db_status_lbl = tk.Label(db_row,text="",fg='#aaaaaa',bg=BG_COLOR,font=SMALL_FONT)
        self.db_status_lbl.grid(row=0,column=6,padx=(12,0))

        # --- Save button ---
        save_frame = tk.Frame(right, bg=BG_COLOR)
        save_frame.pack(fill='x', pady=(8, 0))
        self.save_word_btn = tk.Button(
            save_frame, text="💾 回存WORD檔 / 入資料庫",
            command=self._save_and_archive,
            bg='#cc6600', fg=FG_COLOR, font=DEFAULT_FONT, padx=24, pady=10,
            state='disabled'
        )
        self.save_word_btn.pack(fill='x')

        # Phase 6: 庫存按鈕
        btn_row = tk.Frame(right, bg=BG_COLOR)
        btn_row.pack(fill='x', pady=(4,0))
        tk.Button(btn_row, text="📊 庫存概覽", command=self._show_inventory_panel,
                  bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT, padx=16, pady=6).pack(side='left', padx=4)
        tk.Button(btn_row, text="📋 品項管理", command=self._show_item_manager,
                  bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT, padx=16, pady=6).pack(side='left', padx=4)
        tk.Button(btn_row, text="⚙️ 規則管理", command=self._show_rule_manager,
                  bg=BTN_BG, fg=FG_COLOR, font=SMALL_FONT, padx=16, pady=6).pack(side='left', padx=4)

        self._db_initialized = False

        # Internal state for editing
        self._edit_entry = None
        self._edit_row_id = None
        self._edit_col_idx = None
        self._current_docx_path = None
        self._dirty = False

        mid.add(right, minsize=500)

    def _build_bottom_log(self):
        bottom = tk.Frame(self.root, bg=BG_COLOR)
        bottom.pack(side='bottom', fill='both', expand=False, padx=20, pady=(0, 12))
        tk.Label(bottom, text="處理日誌", fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT).pack(anchor='w')
        self.log_area = scrolledtext.ScrolledText(
            bottom, height=8, bg=LOG_BG, fg=LOG_FG,
            font=MONO_FONT, state='disabled'
        )
        self.log_area.pack(fill='both', expand=True)

    # -------------------
    # Image preview
    # -------------------
    def select_files(self):
        paths = filedialog.askopenfilenames(
            title="選擇圖片檔案（可多選）",
            filetypes=[("Image files", "*.png *.PNG *.jpg *.JPG *.jpeg *.JPEG *.bmp *.BMP *.tiff *.TIFF *.webp *.WEBP")],
            initialdir=r"/mnt/e/DiskCUse/HFDownloads/OCRUse02"
        )
        if not paths:
            return

        new_paths = []
        dup_count = 0
        for p in paths:
            file_hash = self._calc_file_hash(p)
            if file_hash in self.image_hashes:
                dup_count += 1
                self.log(f"跳過重複檔案（已存在雜湊）：{os.path.basename(p)}")
                continue
            # 記錄雜湊
            self.image_hashes[file_hash] = p
            new_paths.append(p)

        if dup_count:
            self.log(f"已略過 {dup_count} 個重複檔案")
            self._save_hash_db()

        if not new_paths:
            self.log("所有選擇檔案皆為重複，未新增")
            return

        self.image_paths = list(new_paths)
        self.current_image_idx = 0
        self.latest_docx_paths = {}
        names = [os.path.basename(p) for p in self.image_paths]
        self.file_combo['values'] = names
        self.file_var.set(names[0])
        self.log(f"已選擇 {len(new_paths)} 個新檔案：{', '.join(names)}")
        self.latest_docx_path = None
        self._update_word_buttons()
        self._draw_image_preview()

    def _on_file_selected(self, event=None):
        idx = self.file_combo.current()
        if idx >= 0 and idx < len(self.image_paths):
            self.current_image_idx = idx
            self._draw_image_preview()
            self._update_word_buttons()

    def _current_image_path(self):
        if self.image_paths and 0 <= self.current_image_idx < len(self.image_paths):
            return self.image_paths[self.current_image_idx]
        return None

    def _draw_image_preview(self):
        path = self._current_image_path()
        if not path:
            self.image_canvas.delete("all")
            self.image_canvas.create_text(
                self.image_canvas.winfo_width() / 2,
                self.image_canvas.winfo_height() / 2,
                text='尚未選擇圖片', fill='#888888', font=SMALL_FONT
            )
            self.zoom_label.config(text='100%')
            return
        try:
            from PIL import Image, ImageTk
            self._raw_image = Image.open(path)
            self._fit_to_canvas()
        except Exception as e:
            self.log(f'圖片預覽失敗：{e}')

    def _render_zoomed_image(self):
        if not hasattr(self, '_raw_image') or self._raw_image is None:
            return
        try:
            from PIL import Image, ImageTk
            img = self._raw_image
            scale = self._zoom_scale
            new_w = int(img.width * scale)
            new_h = int(img.height * scale)
            if new_w < 10 or new_h < 10:
                return
            img_resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
            self.preview_photo = ImageTk.PhotoImage(img_resized)
            self.image_canvas.delete("all")
            self.image_canvas.create_image(
                self._img_offset_x, self._img_offset_y,
                image=self.preview_photo, anchor='nw'
            )
            self.zoom_label.config(text=f'{int(scale*100)}%')
        except Exception as e:
            self.log(f'圖片渲染失敗：{e}')

    def _zoom_in(self):
        self._zoom_scale = min(self._zoom_scale * 1.25, 10.0)
        self._render_zoomed_image()

    def _zoom_out(self):
        self._zoom_scale = max(self._zoom_scale / 1.25, 0.1)
        self._render_zoomed_image()

    def _zoom_reset(self):
        self._zoom_scale = 1.0
        self._img_offset_x = 0
        self._img_offset_y = 0
        self._fit_to_canvas()

    def _fit_to_canvas(self):
        if not hasattr(self, '_raw_image') or self._raw_image is None:
            return
        try:
            cw = max(self.image_canvas.winfo_width(), 400)
            ch = max(self.image_canvas.winfo_height(), 300)
            img = self._raw_image
            img_ratio = img.width / img.height
            canvas_ratio = cw / ch
            if img_ratio > canvas_ratio:
                fit_w = cw
                fit_h = int(cw / img_ratio)
            else:
                fit_h = ch
                fit_w = int(ch * img_ratio)
            self._zoom_scale = fit_w / img.width
            self._img_offset_x = 0
            self._img_offset_y = 0
            self._render_zoomed_image()
        except Exception as e:
            self.log(f'自適應失敗：{e}')

    def _on_img_drag_start(self, event):
        self._drag_start_x = event.x
        self._drag_start_y = event.y

    def _on_img_drag_move(self, event):
        if self._drag_start_x is None:
            return
        dx = event.x - self._drag_start_x
        dy = event.y - self._drag_start_y
        self._img_offset_x += dx
        self._img_offset_y += dy
        self._drag_start_x = event.x
        self._drag_start_y = event.y
        self._render_zoomed_image()

    def _on_img_drag_end(self, event):
        self._drag_start_x = None
        self._drag_start_y = None

    def _on_mousewheel(self, event):
        """滾輪縮放：Windows (event.delta) / Linux (event.num)"""
        if not hasattr(self, '_raw_image') or self._raw_image is None:
            return
        # Windows: event.delta 正=向上, 負=向下
        # Linux: event.num 4=向上, 5=向下
        if hasattr(event, 'delta') and event.delta:
            direction = 1 if event.delta > 0 else -1
        elif hasattr(event, 'num') and event.num in (4, 5):
            direction = 1 if event.num == 4 else -1
        else:
            return
        factor = 1.15
        if direction > 0:
            self._zoom_scale = min(self._zoom_scale * factor, 10.0)
        else:
            self._zoom_scale = max(self._zoom_scale / factor, 0.1)
        self._render_zoomed_image()

    # -------------------
    # Output dir
    # -------------------
    def select_out_dir(self):
        path = filedialog.askdirectory(title="選擇輸出目錄", initialdir=OUTPUT_DIR)
        if path:
            self.out_var.set(path)

    # -------------------
    # API manager
    # -------------------
    def open_api_manager(self):
        win = tk.Toplevel(self.root)
        win.title("修改 API Key")
        win.geometry("780x520")
        win.configure(bg=BG_COLOR)

        tk.Label(win, text="模型", fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT, width=12, anchor='w').grid(row=0, column=0, padx=10, pady=10)
        tk.Label(win, text="原始 API Key", fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT, width=14, anchor='w').grid(row=0, column=1, padx=10, pady=10)

        model_names = ["llamaindex", "nanonets", "ocr.space"]
        entries = {}
        for i, name in enumerate(model_names, start=1):
            tk.Label(win, text=name, fg=FG_COLOR, bg=BG_COLOR, font=SMALL_FONT, width=12, anchor='w').grid(row=i, column=0, padx=10, pady=8, sticky='w')
            var = tk.StringVar(value=self.api_keys.get(name, ""))
            e = tk.Entry(win, textvariable=var, width=58, bg=ENTRY_BG, fg=FG_COLOR, font=SMALL_FONT)
            e.grid(row=i, column=1, padx=10, pady=8)
            entries[name] = var

        status = tk.Label(win, text="", fg='orange', bg=BG_COLOR, font=SMALL_FONT)
        status.grid(row=4, column=0, columnspan=2, sticky='w', padx=10)

        def save_all():
            changed = False
            for name, var in entries.items():
                new_key = var.get()
                if self.api_keys.get(name) != new_key:
                    if self._save_key_to_json(name, new_key):
                        status.config(text=f"已更新：{name}", fg='#00ff00')
                        changed = True
                    else:
                        status.config(text=f"更新失敗：{name}", fg='red')
                        return
            if not changed:
                status.config(text="無變更", fg='orange')
            else:
                self.log("API Key 已更新並寫回 WebOcrAPI.json")

        tk.Button(win, text="儲存全部修改", command=save_all, bg='#0066cc', fg=FG_COLOR, font=SMALL_FONT, padx=16, pady=6).grid(row=5, column=1, sticky='e', padx=10, pady=12)

    # -------------------
    # Process
    # -------------------
    def start_process(self):
        if self.running:
            messagebox.showinfo("提示", "正在處理中，請稍候。")
            return
        if not self.image_paths:
            messagebox.showerror("錯誤", "請先選擇圖片檔案")
            return

        out_dir = self.out_var.get().strip()
        if not out_dir:
            out_dir = OUTPUT_DIR
        os.makedirs(out_dir, exist_ok=True)

        self.running = True
        self._batch_results = {}
        # disable start button during batch
        for btn in self._get_all_buttons():
            if btn.winfo_exists():
                btn.config(state='disabled')
        self.log(f"開始批次辨識 {len(self.image_paths)} 張圖片...")
        threading.Thread(target=self._batch_process_worker, args=(out_dir,), daemon=True).start()

    def _get_all_buttons(self):
        """返回所有需要在批次處理期間禁用的按鈕"""
        btns = []
        # top bar buttons
        for child in self.root.winfo_children():
            btns.extend(self._collect_buttons(child))
        return btns

    def _collect_buttons(self, widget):
        btns = []
        if isinstance(widget, tk.Button):
            btns.append(widget)
        elif hasattr(widget, 'winfo_children'):
            for child in widget.winfo_children():
                btns.extend(self._collect_buttons(child))
        return btns

    def _batch_process_worker(self, out_dir: str):
        """逐張辨識：處理完一張 → 立即顯示表格 → 繼續下一張"""
        total = len(self.image_paths)
        for idx, image_path in enumerate(self.image_paths, 1):
            if not self.running:
                break
            self.root.after(0, lambda i=idx, p=image_path: self.log(f"[{i}/{total}] 開始辨識：{os.path.basename(p)}"))
            docx_path = self._process_single_image(out_dir, image_path)
            if docx_path:
                self._batch_results[image_path] = docx_path
                # 立即切換到此圖片，顯示其表格
                self.root.after(0, lambda p=image_path, d=docx_path: self._show_result_for_image(p, d))
            else:
                self.root.after(0, lambda i=idx: self.log(f"[{i}/{total}] 辨識失敗"))
        # 全部完成，恢復按鈕（不彈提示）
        self.root.after(0, self._on_batch_done)

    def _show_result_for_image(self, image_path: str, docx_path: str):
        """切換到指定圖片並顯示其表格"""
        try:
            idx = self.image_paths.index(image_path)
        except ValueError:
            return
        self.current_image_idx = idx
        self.file_var.set(os.path.basename(image_path))
        self._draw_image_preview()
        self._update_word_buttons()

    def _on_batch_done(self):
        """批次完成：恢復按鈕，不提示"""
        self.running = False
        for btn in self._get_all_buttons():
            if btn.winfo_exists():
                btn.config(state='normal')
        self.log(f"批次辨識完成，成功 {len(self._batch_results)}/{len(self.image_paths)}")

    def _process_single_image(self, out_dir: str, image_path: str):
        """處理單張圖片，返回 docx 路徑或 None"""
        errors = []
        last_docx = None
        jpeg_to_clean = None

        # 統一轉 JPEG（JPG/JPEG 直接跳過）
        original_path = image_path
        converted = convert_to_jpeg(image_path)
        if converted != image_path:
            jpeg_to_clean = converted
            self.log(f"格式轉換：{os.path.basename(image_path)} → {os.path.basename(converted)}")
        image_path = converted

        # 1) LlamaIndex
        try:
            self.log(f"[核心] LlamaIndex 開始辨識：{os.path.basename(image_path)} ...")
            res = llamaindex_process(
                api_key=self.api_keys.get("llamaindex", ""),
                image_path=image_path,
                output_dir=out_dir,
                save_word=True,
                save_excel=False,
                save_html=False,
                save_md=False,
            )
            last_docx = res.get("word")
            self.log(f"[核心] LlamaIndex 完成：{res}")
        except Exception as e:
            msg = str(e)
            errors.append(f"LlamaIndex 失敗：{msg}")
            self.log(f"[核心] LlamaIndex 失敗：{msg}")

        # 2) Nanonets
        if not last_docx:
            try:
                self.log(f"[核心] Nanonets 開始辨識：{os.path.basename(image_path)} ...")
                res = nanonets_process(
                    api_key=self.api_keys.get("nanonets", ""),
                    image_path=image_path,
                    output_dir=out_dir,
                    save_word=True,
                    save_excel=False,
                    save_md=False,
                )
                last_docx = res.get("word")
                self.log(f"[核心] Nanonets 完成：{res}")
            except Exception as e:
                msg = str(e)
                errors.append(f"Nanonets 失敗：{msg}")
                self.log(f"[核心] Nanonets 失敗：{msg}")

        # 3) OCR.Space
        if not last_docx:
            try:
                self.log(f"[核心] OCR.Space 開始辨識：{os.path.basename(image_path)} ...")
                res = ocrspace_process(
                    api_key=self.api_keys.get("ocr.space", ""),
                    image_path=image_path,
                    output_dir=out_dir,
                    save_word=True,
                    save_excel=False,
                    save_text=False,
                    save_md=False,
                )
                last_docx = res.get("word")
                self.log(f"[核心] OCR.Space 完成：{res}")
            except Exception as e:
                msg = str(e)
                errors.append(f"OCR.Space 失敗：{msg}")
                self.log(f"[核心] OCR.Space 失敗：{msg}")

        # 清理暫存 JPEG
        if jpeg_to_clean and os.path.exists(jpeg_to_clean):
            try:
                os.remove(jpeg_to_clean)
            except OSError:
                pass

        # 回報結果（使用原始路徑）
        if last_docx and os.path.exists(last_docx):
            self.latest_docx_paths[original_path] = last_docx
            self.root.after(0, lambda p=last_docx: self.log(f"完成：{p}"))
            return last_docx
        elif errors:
            self.root.after(0, lambda e=errors: self.log(f"全部引擎失敗：{'; '.join(e)}"))
            return None
        else:
            self.root.after(0, lambda: self.log("已處理，但未產出 Word"))
            return None

    

    def _process_worker(self, out_dir: str, image_path: str):
        """保留舊方法以防相容，不再使用"""
        pass

    # -------------------
    # Word actions — LibreOffice Writer
    # -------------------
    def _update_word_buttons(self):
        """Enable/disable buttons and load table based on current image's docx status."""
        image = self._current_image_path()
        docx = self.latest_docx_paths.get(image) if image else None
        has_docx = docx and os.path.exists(docx)
        if docx:
            # 只顯示檔名，字體 12
            self.word_path_var.set(os.path.basename(docx))
        elif image:
            self.word_path_var.set("（此圖片尚未產出 Word）")
        else:
            self.word_path_var.set("")
        if has_docx:
            self.save_word_btn.config(state='normal')
            self.word_status.config(text="✅ 已就緒", fg='#00ff00')
            # Load table into Treeview if not already loaded for this docx
            if self._current_docx_path != docx:
                self._load_docx_to_treeview(docx)
        else:
            self.save_word_btn.config(state='disabled')
            # Clear tree
            self.tree.delete(*self.tree.get_children())
            self.tree['columns'] = []
            self._clear_after_table_text()
            self._current_docx_path = None
            self._dirty = False
            self.after_text_cache = []
            if self.running:
                self.word_status.config(text="⏳ 辨識中...", fg='orange')
            else:
                self.word_status.config(text="（尚未產出）", fg='orange')

    def _clear_after_table_text(self):
        """清除表格下方文字區域"""
        self.after_table_text.config(state='normal')
        self.after_table_text.delete('1.0', tk.END)
        # 保持 normal 以便使用者編輯（下方文字含日期/下收手寫字）
        # 同步清除表格上方文字區域
        if hasattr(self, 'before_table_text'):
            self.before_table_text.config(state='normal')
            self.before_table_text.delete('1.0', tk.END)

    # -------------------
    # Word table editor — Treeview
    # -------------------
    def _load_docx_to_treeview(self, docx_path: str):
        """載入 first table + 表格後文字段落"""
        try:
            from docx import Document
            doc = Document(docx_path)
            if not doc.tables:
                self.log("Word 檔案中無表格")
                return
            table = doc.tables[0]

            # Extract headers
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            if not headers:
                self.log("表格無表頭")
                return

            # Configure tree columns
            self.tree['columns'] = headers
            for h in headers:
                self.tree.heading(h, text=h, anchor='w')
                self.tree.column(h, width=120, minwidth=80, stretch=True)

            # Clear existing rows
            self.tree.delete(*self.tree.get_children())

            # Insert data rows (skip header row 0)
            for row_idx, row in enumerate(table.rows[1:], start=1):
                values = [cell.text.strip() for cell in row.cells]
                if len(values) < len(headers):
                    values += [''] * (len(headers) - len(values))
                elif len(values) > len(headers):
                    values = values[:len(headers)]
                self.tree.insert('', 'end', iid=f'row_{row_idx}', values=values)

            # 擷取表格後文字段落（表格之後的所有 paragraphs）
            after_text = []
            # 擷取表格前文字段落（表格上方：含庫別/日期標頭資訊）
            before_text = []
            table_end_element = table._element
            found_table = False
            for para in doc.element.body:
                if para is table_end_element:
                    found_table = True
                    continue
                if found_table:
                    # 擷取後續段落文字（跳過其他表格元素）
                    from docx.oxml.ns import qn
                    if para.tag.endswith('}p'):
                        text = para.text.strip() if hasattr(para, 'text') else ''
                        if text:
                            after_text.append(text)
                else:
                    # 表格之前的段落 = 上方文字
                    if para.tag.endswith('}p'):
                        text = para.text.strip() if hasattr(para, 'text') else ''
                        if text:
                            before_text.append(text)
            # 顯示下方文字
            self.after_table_text.config(state='normal')
            self.after_table_text.delete('1.0', tk.END)
            if after_text:
                self.after_table_text.insert('1.0', '\n'.join(after_text))
                self.after_table_label.pack()
                self.after_table_frame.pack(fill='x', pady=(4, 0))
            else:
                self.after_table_text.insert('1.0', '（無）')
                self.after_table_label.pack()
                self.after_table_frame.pack(fill='x', pady=(4, 0))
            self.after_table_text.config(state='normal')

            # 顯示上方文字 (供日期/庫別辨識錯誤時修改)
            self.before_table_text.config(state='normal')
            self.before_table_text.delete('1.0', tk.END)
            if before_text:
                self.before_table_text.insert('1.0', '\n'.join(before_text))
                self.before_table_label.pack()
                self.before_table_frame.pack(fill='x', pady=(4, 0))
            else:
                self.before_table_text.insert('1.0', '（無）')
                self.before_table_label.pack()
                self.before_table_frame.pack(fill='x', pady=(4, 0))
            self.before_table_text.config(state='normal')

            self._current_docx_path = docx_path
            self._dirty = False
            self.after_text_cache = after_text  # 保存以備回存
            self.before_text_cache = before_text  # 保存以備回存
            # 自動帶入：日期(表格上方優先) + 庫別(表格上方辨識文字，如 4-1庫)
            # 未回存前自動帶入；使用者手動修改後以修改為準。
            self._auto_fill_from_ocr(before_text, after_text, headers)
            self.log(f"已載入表格：{len(table.rows)-1} 列 × {len(headers)} 欄")
        except Exception as e:
            self.log(f"載入 Word 表格失敗：{e}")
            messagebox.showerror("錯誤", f"無法讀取 Word 表格：{e}")

    def _on_tree_click(self, event):
        """Close any open editor when clicking elsewhere."""
        if self._edit_entry:
            self._finish_edit()

    def _on_cell_double_click(self, event):
        """Start editing the double-clicked cell."""
        if self._edit_entry:
            self._finish_edit()

        region = self.tree.identify('region', event.x, event.y)
        if region != 'cell':
            return

        row_id = self.tree.identify_row(event.y)
        col_id = self.tree.identify_column(event.x)
        if not row_id or not col_id:
            return

        col_idx = int(col_id.replace('#', '')) - 1
        if col_idx < 0 or col_idx >= len(self.tree['columns']):
            return

        # Get cell bounding box
        bbox = self.tree.bbox(row_id, col_id)
        if not bbox:
            return
        x, y, w, h = bbox

        # Get current value
        current_val = self.tree.item(row_id, 'values')[col_idx]

        # Create entry widget on top of cell
        entry = tk.Entry(
            self.tree, font=MONO_FONT,
            bg=ENTRY_BG, fg=FG_COLOR, insertbackground=FG_COLOR,
            borderwidth=1, relief='solid'
        )
        entry.place(x=x, y=y, width=w, height=h)
        entry.insert(0, current_val)
        entry.select_range(0, tk.END)
        entry.focus_set()

        def on_focus_out(e):
            self._finish_edit()

        def on_return(e):
            self._finish_edit()

        entry.bind('<FocusOut>', on_focus_out)
        entry.bind('<Return>', on_return)
        entry.bind('<Escape>', lambda e: self._cancel_edit())

        self._edit_entry = entry
        self._edit_row_id = row_id
        self._edit_col_idx = col_idx

    def _finish_edit(self):
        """Commit the edited value to Treeview."""
        if not self._edit_entry:
            return
        new_val = self._edit_entry.get()
        self._edit_entry.destroy()
        self._edit_entry = None

        if self._edit_row_id and self._edit_col_idx is not None:
            values = list(self.tree.item(self._edit_row_id, 'values'))
            if 0 <= self._edit_col_idx < len(values):
                values[self._edit_col_idx] = new_val
                self.tree.item(self._edit_row_id, values=values)
                self._dirty = True

        self._edit_row_id = None
        self._edit_col_idx = None

    def _cancel_edit(self):
        """Discard edit."""
        if self._edit_entry:
            self._edit_entry.destroy()
            self._edit_entry = None
            self._edit_row_id = None
            self._edit_col_idx = None

    def _auto_fill_from_ocr(self, before_text, after_text, headers):
        """從 OCR 結果自動帶入「表格日期」與「庫別」(未回存前)。
        - 日期：表格上方文字優先，其次下方文字，其次表格日期欄。
        - 庫別：從表格上方文字擷取含「庫/倉」的辨識文字 (如 4-1庫、A倉)。
        使用者手動修改後不會被覆蓋 (僅在此載入時帶入)。
        """
        import re
        # 日期來源：上方 > 下方 > 表格日期欄
        date_sources = (list(before_text) if before_text else []) + (list(after_text) if after_text else [])
        try:
            for ci, h in enumerate(headers):
                if '日期' in str(h) or 'date' in str(h).lower():
                    for rid in self.tree.get_children():
                        vs = self.tree.item(rid, 'values')
                        if ci < len(vs) and str(vs[ci]).strip():
                            date_sources.append(str(vs[ci]))
                            break
        except Exception:
            pass
        for txt in date_sources:
            roc = self._parse_date_to_roc(txt)
            if roc:
                self.biz_date_var.set(roc)
                self.log(f"已自動帶入表格日期：{roc}")
                break

        # 庫別來源：表格上方文字 (含 庫/倉 的行)
        if before_text:
            for txt in before_text:
                lib = self._parse_library_from_text(txt)
                if lib:
                    self.lib_var.set(lib)
                    self.log(f"已自動帶入庫別：{lib}")
                    break

    @staticmethod
    def _parse_library_from_text(text: str) -> str | None:
        """從文字擷取庫別名稱。優先取括號內含「庫/倉」的內容 (如 (4-2庫) -> 4-2庫)。
        若無括號則退回整行含 庫/倉 的片段。
        """
        import re
        s = str(text).strip()
        if not s:
            return None
        # 優先：括號內含 庫/倉 (如 (4-2庫) / (A倉))
        m = re.search(r'[（(]([^（）()]*[庫倉][^（）()]*)[）)]', s)
        if m:
            return re.sub(r'\s+', '', m.group(1))
        # 退回：整行含 庫/倉
        if '庫' in s or '倉' in s:
            return re.sub(r'\s+', '', s)
        return None

    @staticmethod
    def _parse_date_to_iso(text: str) -> str | None:
        """解析文字中的日期為西元 ISO (YYYY-MM-DD)，供資料庫儲存。
        支援：西元 2026-07-22 / 2026/7/22；民國 115/07/22 / 115-07-22；中文 115年7月17日。
        """
        import re
        s = str(text)
        # 西元：2026-07-22 / 2026/7/22 / 2026.7.22
        m = re.search(r'(19|20)(\d{2})[./\-年](\d{1,2})[./\-月](\d{1,2})日?', s)
        if m:
            y = int(m.group(1) + m.group(2)); mo = int(m.group(3)); d = int(m.group(4))
            if 1 <= mo <= 12 and 1 <= d <= 31:
                return f"{y:04d}-{mo:02d}-{d:02d}"
        # 民國：115/07/22 / 115-07-22 / 115年7月17日 (民國年 1~199)
        m = re.search(r'(?<![\d年])(\d{2,3})[./\-年](\d{1,2})[./\-月](\d{1,2})日?', s)
        if m:
            roc = int(m.group(1)); mo = int(m.group(2)); d = int(m.group(3))
            if 1 <= roc <= 199 and 1 <= mo <= 12 and 1 <= d <= 31:
                return f"{roc + 1911:04d}-{mo:02d}-{d:02d}"
        return None

    @staticmethod
    def _parse_date_to_roc(text: str) -> str | None:
        """解析文字中的日期為民國顯示格式 (ROC-MM-DD，月日不足2位補0)。
        例：115年7月17日 -> 115-07-17；115/07/22 -> 115-07-22；2026-07-22 -> 115-07-22。
        """
        import re
        s = str(text)
        # 西元 -> 轉民國
        m = re.search(r'(19|20)(\d{2})[./\-年](\d{1,2})[./\-月](\d{1,2})日?', s)
        if m:
            y = int(m.group(1) + m.group(2)); mo = int(m.group(3)); d = int(m.group(4))
            if 1 <= mo <= 12 and 1 <= d <= 31:
                return f"{y - 1911:03d}-{mo:02d}-{d:02d}"
        # 民國 -> 直接格式化
        m = re.search(r'(?<![\d年])(\d{2,3})[./\-年](\d{1,2})[./\-月](\d{1,2})日?', s)
        if m:
            roc = int(m.group(1)); mo = int(m.group(2)); d = int(m.group(3))
            if 1 <= roc <= 199 and 1 <= mo <= 12 and 1 <= d <= 31:
                return f"{roc:03d}-{mo:02d}-{d:02d}"
        return None

    @staticmethod
    def _normalize_date_to_iso(display: str) -> str:
        """將表格日期欄的顯示值 (可能為民國 115-07-17 或西元 2026-07-17) 轉西元 ISO 供入庫。"""
        d = str(display).strip()
        if not d:
            return ""
        iso = UnifiedOCRApp._parse_date_to_iso(d)
        if iso:
            return iso
        # 已是西元格式則直接回傳
        import re
        if re.match(r'^\d{4}-\d{2}-\d{2}$', d):
            return d
        return d

    def _write_after_text_to_doc(self, doc, table):
        """將 UI「表格下方文字」(含日期/下收手寫字) 寫回 Word 表格後的段落。
        策略：以編輯後文字行取代表格後既有的非空段落；多出的行則新增段落。
        「（無）」佔位字串視為空、不寫回。
        """
        edited = self.after_table_text.get('1.0', tk.END).rstrip('\n')
        new_lines = [ln for ln in edited.split('\n')] if edited and edited != '（無）' else []

        # 收集表格後的段落元素 (p)
        table_el = table._element
        after_paras = []
        found = False
        for el in doc.element.body:
            if el is table_el:
                found = True
                continue
            if found and el.tag.endswith('}p'):
                after_paras.append(el)

        from docx.text.paragraph import Paragraph
        # 逐行覆寫既有段落
        for i, line in enumerate(new_lines):
            if i < len(after_paras):
                Paragraph(after_paras[i], doc).text = line
            else:
                doc.add_paragraph(line)
        # 若編輯後行數變少，將多餘的舊段落清空
        for j in range(len(new_lines), len(after_paras)):
            Paragraph(after_paras[j], doc).text = ''

    def _write_before_text_to_doc(self, doc, table):
        """將 UI「表格上方文字」(含庫別/日期標頭) 寫回 Word 表格前的段落。
        邏輯同 _write_after_text_to_doc，但作用於表格前段落。
        """
        edited = self.before_table_text.get('1.0', tk.END).rstrip('\n')
        new_lines = [ln for ln in edited.split('\n')] if edited and edited != '（無）' else []

        # 收集表格前的段落元素 (p)
        table_el = table._element
        before_paras = []
        found = False
        for el in doc.element.body:
            if el is table_el:
                found = True
                break
            if el.tag.endswith('}p'):
                before_paras.append(el)

        from docx.text.paragraph import Paragraph
        for i, line in enumerate(new_lines):
            if i < len(before_paras):
                Paragraph(before_paras[i], doc).text = line
            else:
                # 表格前新增段落 (插在文件最前)
                doc.element.body.insert(0, doc.add_paragraph(line)._element)
        for j in range(len(new_lines), len(before_paras)):
            Paragraph(before_paras[j], doc).text = ''

    def _save_treeview_to_docx(self):
        """儲存 Treeview 資料回 Word → 保留表格後文字 → 靜默二次確認"""
        if self._edit_entry:
            self._finish_edit()
        if not self._current_docx_path or not os.path.exists(self._current_docx_path):
            self.log("找不到原始 Word 檔案，跳過儲存")
            return
        if not self._dirty:
            self.log("資料未變更，無需儲存")
            return

        try:
            from docx import Document
            doc = Document(self._current_docx_path)
            if not doc.tables:
                self.log("Word 檔案中無表格，跳過儲存")
                return
            table = doc.tables[0]

            headers = list(self.tree['columns'])
            row_ids = self.tree.get_children()

            # Update each data row
            for i, row_id in enumerate(row_ids):
                values = self.tree.item(row_id, 'values')
                word_row_idx = i + 1
                if word_row_idx >= len(table.rows):
                    break
                row = table.rows[word_row_idx]
                for j, val in enumerate(values):
                    if j < len(row.cells):
                        row.cells[j].text = str(val)

            # 保留/更新表格後文字（含日期、下收手寫字）
            self._write_after_text_to_doc(doc, table)
            # 保留/更新表格前文字（含庫別、日期標頭，供修改後寫回）
            self._write_before_text_to_doc(doc, table)
            doc.save(self._current_docx_path)
            self._dirty = False
            self.log(f"已儲存回 Word：{os.path.basename(self._current_docx_path)}")

            # 靜默二次確認：重新讀取 Word，比對表格資料
            self._silent_verify_and_retry(row_ids, headers)

        except Exception as e:
            self.log(f"儲存 Word 失敗：{e}")

    def _silent_verify_and_retry(self, row_ids, headers):
        """靜默比對 Word 表格與 Treeview 資料；不一致則再儲存（最多 3 次）"""
        max_retry = 3
        for attempt in range(1, max_retry + 1):
            try:
                from docx import Document
                doc2 = Document(self._current_docx_path)
                if not doc2.tables:
                    return
                table2 = doc2.tables[0]

                mismatch = False
                mismatches = []
                for i, row_id in enumerate(row_ids):
                    tv_values = self.tree.item(row_id, 'values')
                    word_row_idx = i + 1
                    if word_row_idx >= len(table2.rows):
                        break
                    row = table2.rows[word_row_idx]
                    for j, tv_val in enumerate(tv_values):
                        w_val = row.cells[j].text.strip() if j < len(row.cells) else ''
                        if str(tv_val) != w_val:
                            mismatch = True
                            mismatches.append(f"列{i+1}欄{j}: Treeview={tv_val} Word={w_val}")

                if not mismatch:
                    self.log(f"驗證通過 ✓（第{attempt}次檢查）")
                    return

                self.log(f"驗證不一致（第{attempt}次），重新儲存...")
                # 重新寫入
                for i, row_id in enumerate(row_ids):
                    tv_values = self.tree.item(row_id, 'values')
                    word_row_idx = i + 1
                    if word_row_idx >= len(table2.rows):
                        break
                    row = table2.rows[word_row_idx]
                    for j, val in enumerate(tv_values):
                        if j < len(row.cells):
                            row.cells[j].text = str(val)
                doc2.save(self._current_docx_path)

            except Exception as e:
                self.log(f"驗證/重存失敗（第{attempt}次）：{e}")
                return

        self.log(f"驗證未通過（已重試{max_retry}次）")

    # ================================================================
    # Phase 3+6+7+8: 入庫、庫存面板、品項管理、規則管理
    # ================================================================
    def _ensure_db(self):
        """延遲初始化 DB (Phase 3)"""
        if getattr(self, '_db_initialized', False):
            return True
        try:
            from core.db.repository import init_db, DB_PATH
            from core.ocr_review_service import OCRReviewService
            from core.inventory_service import InventoryService
            init_db()
            self._ocr_service = OCRReviewService()
            self._inv_service = InventoryService(self._ocr_service.repo)
            self._db_initialized = True
            self.log("資料庫已初始化")
            return True
        except Exception as e:
            self.log(f"資料庫初始化失敗：{e}")
            messagebox.showerror("錯誤", f"資料庫初始化失敗：{e}")
            return False

    def _save_and_archive(self):
        """Phase 3/9: 回存 Word + 入資料庫。
        Phase 9: 多品項日結表 → 逐列迴圈入庫；每列各自品項 + 數量欄。
        若表頭為單品項格式（無 進貨量/出貨量 等欄），退回舊單品項邏輯相容。
        """
        self._save_treeview_to_docx()
        if not self._ensure_db():
            return
        from core.db.repository import compute_image_hash

        biz = self.biz_date_var.get().strip()
        lib = self.lib_var.get().strip()
        if not lib:
            self.db_status_lbl.config(text="⚠️ 庫別為空", fg='orange')
            return
        # 庫別為 OCR 辨識文字 (如 4-1庫)，入庫前確保存在 libraries 表 (FK 相容)
        try:
            self._ocr_service.repo.ensure_library(lib, lib_type='storage')
        except Exception as e:
            self.log(f"庫別寫入 libraries 失敗：{e}")
        wp = self._current_docx_path or ""
        img_path = self.image_paths[self.current_image_idx] if self.current_image_idx < len(self.image_paths) else ""
        img_hash = compute_image_hash(Path(img_path)) if img_path else "no_image"

        headers = list(self.tree['columns'])
        row_ids = self.tree.get_children()
        if not row_ids:
            self.db_status_lbl.config(text="⚠️ 表格無資料", fg='orange')
            return

        # 欄位對應 (表頭含關鍵字 -> 資料庫欄位)
        def col_idx(*keys):
            for i, h in enumerate(headers):
                hs = str(h)
                if any(k in hs for k in keys):
                    return i
            return -1

        i_name = col_idx("品項", "item")
        i_prev = col_idx("前日", "期初", "prev")
        i_out = col_idx("出貨", "out")
        i_in = col_idx("進貨", "inbound", "in_")
        i_close = col_idx("庫存", "結存", "closing")
        i_price = col_idx("進價", "price", "單價")
        i_loss = col_idx("損耗", "loss")
        i_note = col_idx("備註", "note", "remark")
        i_qty = col_idx("數量", "quantity", "qty")

        is_multi = i_in >= 0 or i_out >= 0  # 含進貨/出貨欄 => 多品項日結表

        def to_float(v):
            try:
                return float(str(v).strip())
            except (ValueError, TypeError):
                return 0.0

        if is_multi:
            # Phase 9: 逐列多品項入庫
            rows = []
            for rid in row_ids:
                vals = list(self.tree.item(rid, 'values'))
                name = str(vals[i_name]).strip() if i_name >= 0 and i_name < len(vals) else ""
                if not name:
                    continue
                md = self._row_to_markdown(vals, headers)
                row = {
                    "review_date": biz, "library": lib, "item_name": name,
                    "ocr_raw_name": name, "ocr_text": md, "word_path": wp,
                    "source_image_path": img_path, "source_image_hash": img_hash,
                    "prev_stock": to_float(vals[i_prev]) if i_prev >= 0 else 0.0,
                    "outbound_qty": to_float(vals[i_out]) if i_out >= 0 else 0.0,
                    "inbound_qty": to_float(vals[i_in]) if i_in >= 0 else 0.0,
                    "closing_qty": to_float(vals[i_close]) if i_close >= 0 else 0.0,
                    "unit_price": str(vals[i_price]).strip() if i_price >= 0 and i_price < len(vals) else "",
                    "loss_qty": to_float(vals[i_loss]) if i_loss >= 0 else 0.0,
                    "quantity": to_float(vals[i_in]) if i_in >= 0 else (to_float(vals[i_qty]) if i_qty >= 0 else 0.0),
                }
                if i_note >= 0 and i_note < len(vals):
                    row["notes"] = str(vals[i_note]).strip()
                rows.append(row)
            if not rows:
                self.db_status_lbl.config(text="⚠️ 無有效品項列", fg='orange')
                return
            results = self._ocr_service.save_reviewed_rows(rows)
            ok = [r for r in results if r.ok]
            failed = [r for r in results if not r.ok]
            self.db_status_lbl.config(
                text=f"✅ 多品項入庫 {len(ok)}/{len(results)} 列 ({lib})",
                fg='#00cc00')
            self.log(f"多品項入庫完成：成功 {len(ok)} 筆")
            for r in failed:
                self.log(f"⚠️ 入庫失敗：{r.message} {r.errors}")
            # 顯示第一個有 alert 的結果
            for r in results:
                if r.alerts:
                    self._show_alerts(r.alerts)
                    break
        else:
            # 舊單品項相容邏輯
            item_name = self.item_name_var.get().strip()
            if not item_name:
                try:
                    if row_ids and self.tree['columns']:
                        vals = self.tree.item(row_ids[0], 'values')
                        if vals and vals[0]:
                            item_name = str(vals[0]).strip()
                except Exception:
                    pass
            if not item_name:
                self.db_status_lbl.config(text="⚠️ 品項為空", fg='orange')
                return
            qty = self._parse_quantity_from_tree()
            ocr_md = self._tree_to_markdown()
            r = self._ocr_service.save_reviewed_item(
                review_date=biz, library=lib, item_name=item_name,
                ocr_raw_name=item_name, ocr_text=ocr_md, word_path=wp,
                quantity=qty, source_image_path=img_path, source_image_hash=img_hash)
            if r.ok:
                self.db_status_lbl.config(text=f"✅ 已入庫 ({lib}/{item_name})", fg='#00cc00')
                self.log(f"入庫成功 {item_name} ({biz}/{lib})")
                if r.alerts:
                    self._show_alerts(r.alerts)
            else:
                self.db_status_lbl.config(text=f"❌ {r.errors[0][:30]}", fg='red')
                self.log(f"入庫失敗 {r.errors}")

    def _parse_quantity_from_tree(self) -> float:
        try:
            hdrs = list(self.tree['columns'])
            for ci, h in enumerate(hdrs):
                if any(kw in str(h).lower() for kw in ['數量','qty','weight','重量']):
                    total = 0.0
                    for rid in self.tree.get_children():
                        vs = self.tree.item(rid,'values')
                        if ci < len(vs):
                            try: total += float(str(vs[ci]).replace(',',''))
                            except: pass
                    return total
        except: pass
        return 0.0

    def _row_to_markdown(self, vals, headers) -> str:
        """Phase 9: 將單列值轉為 markdown 表格文字（含表頭）。"""
        lines = []
        if headers:
            lines.append("|" + "|".join(str(h) for h in headers) + "|")
            lines.append("|" + "|".join("---" for _ in headers) + "|")
        lines.append("|" + "|".join(str(v).replace('\n', ' ') for v in vals) + "|")
        return "\n".join(lines)

    def _tree_to_markdown(self) -> str:
        lines = []
        hdrs = list(self.tree['columns'])
        if hdrs:
            lines.append("|" + "|".join(str(h) for h in hdrs) + "|")
            lines.append("|" + "|".join("---" for _ in hdrs) + "|")
        for rid in self.tree.get_children():
            V = self.tree.item(rid,'values')
            lines.append("|" + "|".join(str(v).replace('\n',' ') for v in V) + "|")
        return "\n".join(lines)

    def _show_alerts(self, alerts):
        """提醒視窗"""
        dlg = tk.Toplevel(self.root)
        dlg.title(f"⚠️ 提醒 ({len(alerts)})")
        dlg.configure(bg=BG_COLOR); dlg.geometry("600x400")
        cvs = tk.Canvas(dlg,bg=BG_COLOR,highlightthickness=0)
        sb = ttk.Scrollbar(dlg,orient='vertical',command=cvs.yview)
        fr = tk.Frame(cvs,bg=BG_COLOR)
        fr.bind("<Configure>",lambda e:cvs.configure(scrollregion=cvs.bbox("all")))
        cvs.create_window((0,0),window=fr,anchor='nw'); cvs.configure(yscrollcommand=sb.set)
        cl = {'info':'#4488ff','warning':'#ffaa00','critical':'#ff4444'}
        for a in alerts:
            c = cl.get(a.severity,'#ffaa00')
            tk.Label(fr,text=f"[{a.severity.upper()}] {a.rule_name}",fg=c,bg=BG_COLOR,
                     font=('微軟正黑體',12,'bold')).pack(anchor='w',padx=10,pady=(8,0))
            tk.Label(fr,text=a.message,fg=FG_COLOR,bg=BG_COLOR,
                     font=('微軟正黑體',11),wraplength=550,justify='left').pack(anchor='w',padx=10,pady=2)
        cvs.pack(side='left',fill='both',expand=True)
        sb.pack(side='right',fill='y')
        tk.Button(dlg,text="✓ 已知悉",command=dlg.destroy,bg=BTN_BG,fg=FG_COLOR,font=DEFAULT_FONT).pack(pady=10)

    # --- Phase 6: 庫存概覽面板 ---
    def _show_inventory_panel(self):
        if not self._ensure_db(): return
        biz = self.biz_date_var.get().strip()
        lib = self.lib_var.get().strip()
        try:
            diffs = self._inv_service.calculate_daily(biz, lib)
        except Exception as e:
            self.log(f"庫存計算失敗：{e}")
            return

        win = tk.Toplevel(self.root)
        win.title(f"庫存概览 {biz}/{lib}"); win.geometry("760x540"); win.configure(bg=BG_COLOR)
        abnormal = [d for d in diffs.values() if d.is_abnormal]
        normal = [d for d in diffs.values() if not d.is_abnormal]
        # 近日進貨：最近一日有進貨的品項數量 (僅該日有進貨者顯示，其餘留空)
        recent_date, recent_inbound = self._inv_service.get_recent_inbound(before_date=biz)
        recent_hint = f"（近日進貨基準日：{recent_date}）" if recent_date else "（無進貨紀錄）"
        lbl = tk.Label(win,text=f"共 {len(diffs)} 品項（🔴異常 {len(abnormal)} / ⚪正常 {len(normal)}）{recent_hint}",
                       fg=FG_COLOR,bg=BG_COLOR,font=TITLE_FONT).pack(anchor='w',padx=14,pady=(10,4))
        cvs = tk.Canvas(win,bg=BG_COLOR,highlightthickness=0)
        sb = ttk.Scrollbar(win,orient='vertical',command=cvs.yview)
        box = tk.Frame(cvs,bg=BG_COLOR)
        box.bind("<Configure>",lambda e:cvs.configure(scrollregion=cvs.bbox("all")))
        cvs.create_window((0,0),window=box,anchor='nw'); cvs.configure(yscrollcommand=sb.set)
        hdr = tk.Frame(box,bg=BG_COLOR)
        for i,txt in enumerate(["品項","近日進貨","理論","實際","損耗量","損耗率"]):
            tk.Label(hdr,text=txt,fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT,width=10).grid(row=0,column=i,padx=2)
        hdr.pack(anchor='w',padx=8,pady=4)
        for color,grp in [('#ff4444',abnormal),('#cccccc',normal)]:
            for d in grp:
                rw = tk.Frame(box,bg=BG_COLOR)
                # 近日進貨：該品項最近一日有進貨才顯示數量，否則留空
                recent_qty = f"{recent_inbound[d.item_name]:.1f}" if d.item_name in recent_inbound else ""
                for j,val in enumerate([d.item_name,recent_qty,f"{d.expected_qty:.1f}",f"{d.actual_qty:.1f}",f"{d.loss_qty:+.2f}",f"{d.loss_pct:.1f}%"]):
                    tk.Label(rw,text=val,fg=color,bg=BG_COLOR,font=SMALL_FONT,width=10).grid(row=0,column=j,padx=2)
                rw.pack(anchor='w',padx=8)
        cvs.pack(side='left',fill='both',expand=True,padx=8,pady=8)
        sb.pack(side='right',fill='y')
        tk.Button(win,text="關閉",command=win.destroy,bg=BTN_BG,fg=FG_COLOR,font=SMALL_FONT).pack(pady=6)

    # --- Phase 7: 品項管理 ---
    def _show_item_manager(self):
        if not self._ensure_db(): return
        win = tk.Toplevel(self.root); win.title("品項管理"); win.geometry("540x480"); win.configure(bg=BG_COLOR)
        tk.Label(win,text="所有品項 (canonical_items)",fg=FG_COLOR,bg=BG_COLOR,font=TITLE_FONT).pack(anchor='w',padx=14,pady=(10,4))
        cvs = tk.Canvas(win,bg=BG_COLOR,highlightthickness=0)
        sb = ttk.Scrollbar(win,orient='vertical',command=cvs.yview)
        box = tk.Frame(cvs,bg=BG_COLOR)
        box.bind("<Configure>",lambda e:cvs.configure(scrollregion=cvs.bbox("all")))
        cvs.create_window((0,0),window=box,anchor='nw'); cvs.configure(yscrollcommand=sb.set)
        items = self._ocr_service.repo.get_all_canonical_items(active_only=False)
        # 使用 grid 讓各欄(品項/狀態/分類/屬性按鈕)垂直對齊
        for r, ci in enumerate(items):
            f = tk.Frame(box,bg=BG_COLOR)
            tk.Label(f,text=ci.canonical_name,fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT,width=16,anchor='w').grid(row=0,column=0,padx=4,sticky='w')
            status = "✓" if ci.is_active else "✗"
            tk.Label(f,text=status,fg='#00cc00' if ci.is_active else '#ff4444',bg=BG_COLOR,font=SMALL_FONT,width=3).grid(row=0,column=1,padx=4)
            tk.Label(f,text=f"{ci.category or ''}",fg='#aaaaaa',bg=BG_COLOR,font=SMALL_FONT,width=8,anchor='w').grid(row=0,column=2,padx=4,sticky='w')
            tk.Button(f,text="屬性",command=lambda n=ci.canonical_name:self._show_item_attrs(n),
                      bg=BTN_BG,fg=FG_COLOR,font=SMALL_FONT,width=6).grid(row=0,column=3,padx=4)
            f.pack(anchor='w',padx=8,pady=2)
        cvs.pack(side='left',fill='both',expand=True,padx=8,pady=8)
        sb.pack(side='right',fill='y')
        tk.Button(win,text="關閉",command=win.destroy,bg=BTN_BG,fg=FG_COLOR,font=SMALL_FONT).pack(pady=6)

    def _show_item_attrs(self, name):
        """Phase 8: 單品項屬性編輯"""
        win = tk.Toplevel(self.root); win.title(f"{name} 的屬性"); win.geometry("540x400"); win.configure(bg=BG_COLOR)
        tk.Label(win,text=f"品項: {name}",fg=FG_COLOR,bg=BG_COLOR,font=TITLE_FONT).pack(anchor='w',padx=14,pady=(10,4))

        # 分類下拉 (存入 canonical_items.category)
        cur_ci = next((c for c in self._ocr_service.repo.get_all_canonical_items(active_only=False)
                       if c.canonical_name == name), None)
        cat_frame = tk.Frame(win,bg=BG_COLOR); cat_frame.pack(fill='x',padx=14,pady=2)
        tk.Label(cat_frame,text="分類:",fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT,width=16,anchor='w').pack(side='left')
        cat_var = tk.StringVar(value=(cur_ci.category if cur_ci and cur_ci.category else ''))
        ttk.Combobox(cat_frame,textvariable=cat_var,values=["水果","蔬菜","肉類","其他"],
                     state='readonly',width=18,font=SMALL_FONT).pack(side='left')

        attrs = self._ocr_service.get_item_attributes(name)
        entries = {}
        # 屬性 key (資料庫欄位) -> 中文標籤
        attr_labels = {
            'shelf_life_days': '保存天數',
            'normal_loss_pct': '正常損耗率(%)',
            'expiry_date': '到期日(YYYY-MM-DD)',
            'unit': '單位',
        }
        for key in ['shelf_life_days','normal_loss_pct','expiry_date','unit']:
            f = tk.Frame(win,bg=BG_COLOR); f.pack(fill='x',padx=14,pady=2)
            tk.Label(f,text=f"{attr_labels[key]}:",fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT,width=16,anchor='w').pack(side='left')
            val = tk.StringVar(value=attrs.get(key,''))
            tk.Entry(f,textvariable=val,bg=ENTRY_BG,fg=FG_COLOR,font=SMALL_FONT,width=20).pack(side='left')
            entries[key] = val
        def save():
            # 儲存分類
            if cat_var.get().strip():
                self._ocr_service.repo.upsert_canonical_item(name, category=cat_var.get().strip())
            for k,v in entries.items():
                if v.get().strip():
                    self._ocr_service.set_item_attribute(name,k,v.get().strip())
            win.destroy()
            self.log(f"已更新 {name} 屬性")
        tk.Button(win,text="儲存",command=save,bg='#cc6600',fg=FG_COLOR,font=DEFAULT_FONT).pack(pady=10)

    # --- Phase 8: 規則管理 ---
    def _show_rule_manager(self):
        win = tk.Toplevel(self.root); win.title("規則管理"); win.geometry("640x480"); win.configure(bg=BG_COLOR)
        tk.Label(win,text="自訂提醒規則",fg=FG_COLOR,bg=BG_COLOR,font=TITLE_FONT).pack(anchor='w',padx=14,pady=(10,4))
        rules = self._ocr_service.repo.get_enabled_alert_rules() if self._ensure_db() else []
        cvs = tk.Canvas(win,bg=BG_COLOR,highlightthickness=0)
        sb = ttk.Scrollbar(win,orient='vertical',command=cvs.yview)
        box = tk.Frame(cvs,bg=BG_COLOR)
        box.bind("<Configure>",lambda e:cvs.configure(scrollregion=cvs.bbox("all")))
        cvs.create_window((0,0),window=box,anchor='nw'); cvs.configure(yscrollcommand=sb.set)
        if not rules:
            tk.Label(box,text="(尚無自訂規則)",fg='#aaaaaa',bg=BG_COLOR,font=SMALL_FONT).pack(pady=10)
        for R in rules:
            f = tk.Frame(box,bg=BG_COLOR); f.pack(fill='x',padx=8,pady=2)
            tk.Label(f,text=f"{R.rule_name} ({R.rule_type})",fg=FG_COLOR,bg=BG_COLOR,font=SMALL_FONT,width=30).pack(side='left')
            tk.Label(f,text=R.severity,fg='#ffaa00',bg=BG_COLOR,font=SMALL_FONT).pack(side='left',padx=4)
        cvs.pack(side='left',fill='both',expand=True,padx=8,pady=8)
        sb.pack(side='right',fill='y')
        tk.Button(win,text="關閉",command=win.destroy,bg=BTN_BG,fg=FG_COLOR,font=SMALL_FONT).pack(pady=6)

    # -------------------
    # Logging
    # -------------------
    def log(self, msg: str):
        line = f"[{time_str()}] {msg}"
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, line + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')


def main():
    root = tk.Tk()
    app = UnifiedOCRApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()