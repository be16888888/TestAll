#!/usr/bin/env python3
"""
Nanonets OCR Core Module
Contains the core OCR processing logic, separated from UI.
"""

import os
import re
import json
import io
from pathlib import Path
from datetime import datetime
from typing import Dict, Tuple, Optional, Any

import requests
import pandas as pd
from PIL import Image
from docx import Document
from docx.shared import Cm


def load_api_key(json_path: str = "WebOcrAPI.json") -> str:
    """Load Nanonets API key from WebOcrAPI.json.

    Args:
        json_path: Path to the JSON configuration file.

    Returns:
        The API key string.

    Raises:
        FileNotFoundError: If the config file is missing.
        ValueError: If the API key is missing or malformed.
    """
    if not os.path.exists(json_path):
        raise FileNotFoundError(f"Configuration file not found: {json_path}")

    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            configs = json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON in {json_path}: {e}")

    if not isinstance(configs, list):
        raise ValueError(f"Expected a list in {json_path}")

    target = None
    for cfg in configs:
        if cfg.get("ModelName") == "nanonets":
            target = cfg
            break

    if target is None:
        raise ValueError("No configuration for model 'nanonets' found in WebOcrAPI.json")

    api_key = target.get("api_key", "").strip()
    if not api_key:
        raise ValueError("API key for 'nanonets' is empty in WebOcrAPI.json")

    return api_key


def convert_to_jpeg(image_path: str) -> str:
    """Convert any non-JPEG image to JPEG format for Nanonets compatibility.

    Args:
        image_path: Path to the input image file.

    Returns:
        Path to the JPEG image (original if already JPEG, otherwise converted).
    """
    path = Path(image_path)
    suffix = path.suffix.lower()
    if suffix in ('.jpeg', '.jpg'):
        return image_path

    jpeg_path = path.with_suffix('.jpeg')
    with Image.open(image_path) as img:
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        img.save(jpeg_path, 'JPEG', quality=95)
    return str(jpeg_path)


def call_nanonets_api(api_key: str, image_path: str) -> Dict[str, Any]:
    """Call Nanonets OCR API (synchronous extraction).

    Args:
        api_key: Nanonets API key.
        image_path: Path to the image file to process.

    Returns:
        JSON response from the API as a dictionary.

    Raises:
        Exception: If the API returns an error or the request fails.
    """
    url = "https://extraction-api.nanonets.com/api/v1/extract/sync"
    headers = {"Authorization": f"Bearer {api_key}"}
    with open(image_path, 'rb') as f:
        files = {"file": (os.path.basename(image_path), f)}
        data = {"output_format": "markdown"}
        response = requests.post(url, headers=headers, files=files, data=data, timeout=120)

    if response.status_code != 200:
        raise Exception(f"API returned error (HTTP {response.status_code}): {response.text}")

    return response.json()


def extract_markdown_content(json_data: Dict[str, Any]) -> str:
    """Extract markdown content (which contains HTML table) from Nanonets API response.

    Args:
        json_data: The JSON response from Nanonets API.

    Returns:
        The markdown/content string.

    Raises:
        ValueError: If markdown content cannot be found.
    """
    try:
        content = json_data.get("result", {}).get("markdown", {}).get("content")
        if content is None:
            content = json_data.get("markdown", {}).get("content")
        if content is None:
            if isinstance(json_data.get("result"), list) and len(json_data["result"]) > 0:
                first = json_data["result"][0]
                content = first.get("markdown", {}).get("content")
        if content is None:
            raise ValueError("Unable to find markdown content in JSON")
        return content
    except Exception as e:
        raise Exception(f"Failed to parse JSON: {e}\nOriginal response: {json_data}")


def save_as_markdown(content: str, image_path: str, output_dir: str) -> str:
    """Save HTML/Markdown content to a .md file.

    Args:
        content: The HTML/content to save.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved markdown file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_filename = f"{base_name}_{timestamp}.md"
    md_path = os.path.join(output_dir, md_filename)

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# OCR Results - {os.path.basename(image_path)}\n\n")
        f.write(f"**Processing Time**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write(content)
    return md_path


def _split_md_around_table(md_text: str) -> tuple[str, str]:
    """Split Markdown text into parts before and after the first table.

    Args:
        md_text: The Markdown text containing a table.

    Returns:
        A tuple (before_text, after_text). If no table found, after_text is empty.
    """
    lines = md_text.splitlines()
    table_start = None
    table_end = None
    in_table = False
    for i, line in enumerate(lines):
        if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
            if not in_table:
                in_table = True
                table_start = i
        else:
            if in_table:
                table_end = i
                break
    if table_start is None:
        return md_text, ""
    before_lines = lines[:table_start]
    after_lines = lines[table_end:] if table_end is not None else []
    before_text = "\n".join(before_lines).strip()
    after_text = "\n".join(after_lines).strip()
    return before_text, after_text


def _generate_word_filename(before_text: str, fallback_base: str, timestamp: str) -> str:
    """Generate a Word filename based on extracted date and warehouse info.

    Args:
        before_text: Text before the table (may contain date and warehouse).
        fallback_base: Base name to use if extraction fails.
        timestamp: Timestamp string for fallback.

    Returns:
        Suggested filename (without directory).
    """
    date_str = None
    warehouse_str = None
    # Date: support Chinese year/month/day with optional spaces
    date_match = re.search(r'(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', before_text)
    if date_match:
        y, m, d = date_match.groups()
        date_str = f"{y}年{m}月{d}日"
    # Warehouse: support patterns like (6庫) or (4-2庫)
    wh_match = re.search(r'\((\d{1,3}(?:-\d{1,2})?)\s*庫\)', before_text)
    if wh_match:
        warehouse_str = f"({wh_match.group(1)}庫)"
    if date_str and warehouse_str:
        return f"{date_str}{warehouse_str}.docx"
    else:
        return f"{fallback_base}_{timestamp}.docx"


def _merge_title_info(lines: list[str]) -> list[str]:
    """Merge header information lines (warehouse, date, filler, unit) into a single line.

    Args:
        lines: List of lines before the table.

    Returns:
        List of lines with header info merged.
    """
    if not lines:
        return lines
    new_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if '庫)' in line and (line.startswith('#') or '(' in line):
            combined = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if any(next_line.startswith(prefix) for prefix in [
                    '**日期：**', '日期：',
                    '**填單者：**', '填單者：',
                    '**單位：**', '單位：'
                ]):
                    combined.append(next_line)
                    j += 1
                else:
                    break
            merged_line = '     '.join(combined)
            new_lines.append(merged_line)
            i = j
        else:
            new_lines.append(line)
            i += 1
    return new_lines


def _extract_markdown_table(text: str) -> str | None:
    """Extract the first Markdown table from text.

    Args:
        text: Input text.

    Returns:
        The markdown table as a string, or None if not found.
    """
    lines = text.splitlines()
    table_lines = []
    in_table = False
    for line in lines:
        if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
            if not in_table:
                in_table = True
            table_lines.append(line)
        else:
            if in_table:
                if len(table_lines) >= 3:
                    break
                else:
                    table_lines = []
                    in_table = False
    if len(table_lines) >= 3:
        return "\n".join(table_lines)
    return None


def _parse_markdown_table(md_text: str) -> pd.DataFrame:
    """Parse a Markdown table string into a DataFrame (fallback implementation).

    Args:
        md_text: Markdown table string.

    Returns:
        pandas DataFrame.
    """
    lines = md_text.splitlines()
    data_lines = []
    for line in lines:
        if re.search(r'^\s*\|\s*[-:|]+\s*\|', line):
            continue
        if '|' in line:
            stripped = line.strip('|')
            cells = [cell.strip() for cell in stripped.split('|')]
            data_lines.append(cells)
    if not data_lines:
        raise ValueError("No valid data rows")
    header = data_lines[0]
    data = data_lines[1:]
    max_cols = max(len(row) for row in data_lines)
    for row in data:
        while len(row) < max_cols:
            row.append('')
    while len(header) < max_cols:
        header.append('')
    return pd.DataFrame(data, columns=header[:max_cols])


def extract_table_to_dataframe(md_content: str) -> pd.DataFrame:
    """Extract the first table from Markdown content and convert to DataFrame.

    Args:
        md_content: Markdown string that may contain a table.

    Returns:
        pandas DataFrame containing the table data.

    Raises:
        ValueError: If no table is found or parsing fails.
    """
    # Try to extract markdown table
    md_table = _extract_markdown_table(md_content)
    if md_table:
        return _parse_markdown_table(md_table)
    raise ValueError("No Markdown table found")


def save_as_word(df: pd.DataFrame, image_path: str, output_dir: str, md_content: str) -> str:
    """Save DataFrame as a Word document, with header info extracted from markdown.

    Args:
        df: DataFrame containing the table data.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.
        md_content: The markdown content (used to extract header info).

    Returns:
        Full path to the saved Word file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Split markdown to get text before and after table for naming and content
    before_text, after_text = _split_md_around_table(md_content)
    word_filename = _generate_word_filename(before_text, base_name, timestamp)
    word_path = os.path.join(output_dir, word_filename)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    if before_text:
        before_lines = [line.strip() for line in before_text.split('\n') if line.strip()]
        merged_lines = _merge_title_info(before_lines)
        for line in merged_lines:
            doc.add_paragraph(line)

    if df.shape[1] > 0:
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'
        for j, col_name in enumerate(df.columns):
            table.cell(0, j).text = str(col_name)
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                cell_text = "" if pd.isna(value) else str(value)
                table.cell(i + 1, j).text = cell_text

    if after_text:
        for line in after_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    doc.save(word_path)
    return word_path


def save_as_excel(df: pd.DataFrame, image_path: str, output_dir: str) -> str:
    """Save DataFrame as an Excel file.

    Args:
        df: DataFrame containing the table data.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved Excel file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    excel_filename = f"{base_name}_{timestamp}.xlsx"
    excel_path = os.path.join(output_dir, excel_filename)
    df.to_excel(excel_path, index=False)
    return excel_path


def process_single(api_key: str, image_path: str, output_dir: str,
                   save_word: bool = True, save_excel: bool = True,
                   save_md: bool = True) -> Dict[str, str]:
    """Process a single image through the full Nanonets OCR pipeline.

    Args:
        api_key: Nanonets API key.
        image_path: Path to the image file to process.
        output_dir: Directory where output files should be saved.
        save_word: Whether to save a Word document.
        save_excel: Whether to save an Excel file.
        save_md: Whether to save a Markdown file.

    Returns:
        Dictionary with keys 'md', 'word', 'excel' (if saved) mapping to file paths.

    Raises:
        Exception: If any step in the processing fails.
    """
    # 1. Format conversion
    processed_path = convert_to_jpeg(image_path)

    try:
        # 2. Call API
        response_json = call_nanonets_api(api_key, processed_path)

        # 3. Save raw JSON (optional, useful for debugging)
        base_name = os.path.splitext(os.path.basename(image_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        json_filename = f"{base_name}_{timestamp}_raw.json"
        json_path = os.path.join(output_dir, json_filename)
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(response_json, f, indent=2, ensure_ascii=False)

        # 4. Extract HTML content
        html_content = extract_markdown_content(response_json)

        results = {}

        # 5. Save Markdown if requested
        if save_md:
            md_path = save_as_markdown(html_content, image_path, output_dir)
            results['md'] = md_path

        # 6. Parse table and save Word/Excel if requested
        df = None
        if save_word or save_excel:
            df = extract_table_to_dataframe(html_content)
            if save_word:
                word_path = save_as_word(df, image_path, output_dir, html_content)
                results['word'] = word_path
            if save_excel:
                excel_path = save_as_excel(df, image_path, output_dir)
                results['excel'] = excel_path

        return results

    except Exception as e:
        raise Exception(f"Processing failed: {e}")
    finally:
        # Clean up temporary JPEG file if we created one
        if processed_path != image_path:
            try:
                os.remove(processed_path)
            except OSError:
                pass  # Ignore cleanup errors