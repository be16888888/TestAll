#!/usr/bin/env python3
"""
Nanonets OCR Service Module
Handles all business logic for interacting with Nanonets OCR API and processing results.
"""

import os
import re
import json
import time
from pathlib import Path
from datetime import datetime
from typing import Dict, Tuple, Optional, Any

import requests
import pandas as pd
from PIL import Image
from docx import Document


def convert_to_jpeg(image_path: str) -> str:
    """
    Convert any non-JPEG image to .jpeg format for Nanonets compatibility.
    
    Args:
        image_path: Path to the input image file
        
    Returns:
        Path to the JPEG image (original if already JPEG, otherwise converted)
    """
    path = Path(image_path)
    suffix = path.suffix.lower()
    if suffix in ('.jpeg', '.jpg'):
        return image_path

    jpeg_path = path.with_suffix('.jpeg')
    
    with Image.open(image_path) as img:
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        img.save(jpeg_path, 'JPEG', quality=95)
    
    return str(jpeg_path)


def call_nanonets_api(api_key: str, image_path: str) -> Dict[str, Any]:
    """
    Call Nanonets OCR API (synchronous extraction).
    
    Args:
        api_key: Nanonets API key
        image_path: Path to the image file to process
        
    Returns:
        JSON response from the API as a dictionary
        
    Raises:
        Exception: If the API returns an error or the request fails
    """
    url = "https://extraction-api.nanonets.com/api/v1/extract/sync"
    headers = {"Authorization": f"Bearer {api_key}"}
    
    with open(image_path, 'rb') as f:
        files = {"file": (os.path.basename(image_path), f)}
        data = {"output_format": "markdown"}
        response = requests.post(url, headers=headers, files=files, data=data, timeout=120)
    
    if response.status_code != 200:
        raise Exception(f"API returned error (HTTP {response.status_code}): {response.text}")
    
    return response.json()


def extract_html_content(json_data: Dict[str, Any]) -> str:
    """
    Extract HTML 
    Extract HTML content from Nanonets API response.
    The actual content is in the markdown.content field (which contains HTML table).

    Args:
        json_data: The JSON response from Nanonets API

    Returns:
        The HTML content string

    Raises:
        ValueError: If markdown content cannot be found in the response
    """
    try:
        content = json_data.get("result", {}).get("markdown", {}).get("content")
        if content is None:
            content = json_data.get("markdown", {}).get("content")
        if content is None:
            if isinstance(json_data.get("result"), list) and len(json_data["result"]) > 0:
                first = json_data["result"][0]
                content = first.get("markdown", {}).get("content")
        if content is None:
            raise ValueError("Unable to find markdown content in JSON")
        return content
    except Exception as e:
        raise Exception(f"Failed to parse JSON: {e}\\nOriginal response: {json_data}")


def html_table_to_dataframe(html_str: str) -> pd.DataFrame:
    """
    Extract the first <table> from an HTML string and convert to DataFrame.

    Args:
        html_str: String containing HTML with at least one <table> element

    Returns:
        pandas DataFrame containing the table data

    Raises:
        ValueError: If no table is found or table cannot be parsed
    """
    table_pattern = re.compile(r'<table.*?>.*?</table>', re.DOTALL | re.IGNORECASE)
    table_matches = table_pattern.findall(html_str)
    if not table_matches:
        raise ValueError("No table tags found")

    first_table = table_matches[0]
    try:
        dfs = pd.read_html(first_table)
        if not dfs:
            raise ValueError("Parsing table yielded no data")
        return dfs[0]
    except Exception as e:
        raise Exception(f"Failed to parse HTML table: {e}\\nTable snippet:\\n{first_table[:500]}")


def save_as_markdown(content: str, image_path: str, output_dir: str) -> str:
    """
    Save HTML/Markdown content to a .md file.

    Args:
        content: The HTML/content to save
        image_path: Path to the original image (for naming)
        output_dir: Directory where the file should be saved

    Returns:
        Full path to the saved markdown file
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_filename = f"{base_name}_{timestamp}.md"
    md_path = os.path.join(output_dir, md_filename)
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# OCR Results - {os.path.basename(image_path)}\\n\\n")
        f.write(f"**Processing Time**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\\n\\n")
        f.write(content)
    
    return md_path


def save_as_word(df: pd.DataFrame, image_path: str, output_dir: str) -> str:
    """
    Save DataFrame as a Word document.

    Args:
        df: DataFrame containing the table data
        image_path: Path to the original image (for naming)
        output_dir: Directory where the file should be saved

    Returns:
        Full path to the saved Word file
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_filename = f"{base_name}_{timestamp}.docx"
    word_path = os.path.join(output_dir, word_filename)
    
    doc = Document()
    doc.add_heading('Extracted Table Data', level=1)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Table Grid'
    for j, col_name in enumerate(df.columns):
        table.cell(0, j).text = str(col_name)
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            table.cell(i + 1, j).text = str(value)
    doc.save(word_path)
    
    return word_path


def save_as_excel(df: pd.DataFrame, image_path: str, output_dir: str) -> str:
    """
    Save DataFrame as an Excel file.

    Args:
        df: DataFrame containing the table data
        image_path: Path to the original image (for naming)
        output_dir: Directory where the file should be saved

    Returns:
        Full path to the saved Excel file
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    excel_filename = f"{base_name}_{timestamp}.xlsx"
    excel_path = os.path.join(output_dir, excel_filename)
    
    df.to_excel(excel_path, index=False)
    
    return excel_path


def process_single(api_key: str, image_path: str, output_dir: str) -> Dict[str, str]:
    """
    Process a single image through the full OCR pipeline.

    Args:
        api_key: Nanonets API key
        image_path: Path to the image file to process
        output_dir: Directory where output files should be saved

    Returns:
        Dictionary with keys 'md', 'word', 'excel' (if saved) mapping to file paths

    Raises:
        Exception: If any step in the processing fails
    """
    # 1. Format conversion
    processed_path = convert_to_jpeg(image_path)
    
    try:
        # 2. Call API
        response_json = call_nanonets_api(api_key, processed_path)
        
        # 3. Save raw JSON (optional, but useful for debugging)
        base_name = os.path.splitext(os.path.basename(image_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        json_filename = f"{base_name}_{timestamp}_raw.json"
        json_path = os.path.join(output_dir, json_filename)
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(response_json, f, indent=2, ensure_ascii=False)
        
        # 4. Extract HTML content
        html_content = extract_html_content(response_json)
        
        results = {}
        
        # 5. Save Markdown if requested
        md_path = save_as_markdown(html_content, image_path, output_dir)
        results['md'] = md_path
        
        # 6. Parse table and save Word/Excel if requested
        df = html_table_to_dataframe(html_content)
        
        word_path = save_as_word(df, image_path, output_dir)
        results['word'] = word_path
        
        excel_path = save_as_excel(df, image_path, output_dir)
        results['excel'] = excel_path
        
        return results
        
    except Exception as e:
        raise Exception(f"Processing failed: {e}")
    finally:
        # Clean up temporary JPEG file if we created one
        if processed_path != image_path:
            try:
                os.remove(processed_path)
            except OSError:
                pass  # Ignore errors in cleanup