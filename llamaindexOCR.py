import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Checkbutton, IntVar
import requests
import pandas as pd
from docx import Document
from docx.shared import Cm
import os
import re
import io
import html
from datetime import datetime
import time
import json

class LlamaCloudOCRApp:
    def __init__(self, root):
        self.root = root
        self.root.title("LlamaCloud OCR 表格提取工具 (v2)")
        self.root.geometry("650x670")
        self.root.configure(bg='black')

        fg_color = 'white'
        bg_color = 'black'
        entry_bg = '#333333'
        btn_bg = '#555555'

        title = tk.Label(root, text="LlamaCloud OCR 表格提取工具 (v2)", font=('Arial', 16, 'bold'),
                         fg=fg_color, bg=bg_color)
        title.pack(pady=10)

        # API Key 輸入
        frame_key = tk.Frame(root, bg=bg_color)
        frame_key.pack(pady=5, fill='x', padx=20)
        tk.Label(frame_key, text="API Key：", fg=fg_color, bg=bg_color,
                 font=('Arial', 10)).pack(side='left')
        self.api_entry = tk.Entry(frame_key, width=50, bg=entry_bg, fg=fg_color,
                                  insertbackground='white')
        self.api_entry.pack(side='left', padx=5)

        # 輸出目錄選擇
        frame_out = tk.Frame(root, bg=bg_color)
        frame_out.pack(pady=5, fill='x', padx=20)
        tk.Label(frame_out, text="輸出目錄：", fg=fg_color, bg=bg_color,
                 font=('Arial', 10)).pack(side='left')
        self.dir_entry = tk.Entry(frame_out, width=40, bg=entry_bg, fg=fg_color)
        self.dir_entry.pack(side='left', padx=5)
        self.dir_entry.insert(0, r"E:\DiskCUse\HFDownloads")
        tk.Button(frame_out, text="瀏覽", command=self.select_dir,
                  bg=btn_bg, fg=fg_color).pack(side='left', padx=5)

        # 圖片選擇
        frame_file = tk.Frame(root, bg=bg_color)
        frame_file.pack(pady=10, fill='x', padx=20)
        self.file_label = tk.Label(frame_file, text="未選擇圖片", fg='gray', bg=bg_color,
                                   width=40, anchor='w')
        self.file_label.pack(side='left')
        tk.Button(frame_file, text="選擇圖片", command=self.select_file,
                  bg=btn_bg, fg=fg_color, padx=10).pack(side='right')

        # 輸出格式勾選框
        frame_options = tk.LabelFrame(root, text="儲存格式選項", bg=bg_color, fg=fg_color,
                                      font=('Arial', 10))
        frame_options.pack(pady=10, fill='x', padx=20)

        self.var_word = IntVar(value=1)
        self.var_excel = IntVar(value=0)
        self.var_html = IntVar(value=1)
        self.var_md = IntVar(value=1)

        chk_word = Checkbutton(frame_options, text="儲存 WORD 檔", variable=self.var_word,
                               bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_word.pack(side='left', padx=10)

        chk_excel = Checkbutton(frame_options, text="儲存 EXCEL 檔", variable=self.var_excel,
                                bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_excel.pack(side='left', padx=10)

        chk_html = Checkbutton(frame_options, text="儲存原始 HTML 內容", variable=self.var_html,
                               bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_html.pack(side='left', padx=10)

        chk_md = Checkbutton(frame_options, text="儲存 MD 文件", variable=self.var_md,
                             bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_md.pack(side='left', padx=10)

        # 開始按鈕
        self.process_btn = tk.Button(root, text="🚀 開始辨識並轉換", command=self.process,
                                     bg='#0066cc', fg=fg_color, font=('Arial', 12, 'bold'),
                                     padx=20, pady=5)
        self.process_btn.pack(pady=15)

        # 日誌區域
        self.log_area = scrolledtext.ScrolledText(root, height=18, bg='#1a1a1a',
                                                  fg='#00ff00', font=('Consolas', 10))
        self.log_area.pack(fill='both', expand=True, padx=20, pady=10)
        self.log_area.insert(tk.END, "就緒，請填寫 API Key 並選擇圖片。\n")
        self.log_area.config(state='disabled')

        self.image_path = None

        # 自動載入（或建立）WebOcrAPI.json
        self.ensure_and_load_json()

    def ensure_and_load_json(self):
        """確保 WebOcrAPI.json 存在，並載入 API Key"""
        json_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "WebOcrAPI.json")
        
        if not os.path.exists(json_path):
            default_config = {
                "url": "https://api.cloud.llamaindex.ai",
                "base_url": "https://api.cloud.llamaindex.ai/api/v2",
                "api_key": "llx-kp6L7gAeSDVvtfS4ZXXwkFeQNG82KfcoSxpoC1Wqga1cr5Xa"
            }
            try:
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(default_config, f, indent=2, ensure_ascii=False)
                self.log("已自動建立 WebOcrAPI.json（內含預設 API Key）")
            except Exception as e:
                self.log(f"建立 WebOcrAPI.json 失敗：{e}")
                return

        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            api_key = config.get("api_key") or config.get("LLAMA_CLOUD_API_KEY") or config.get("API_KEY")
            if api_key:
                self.api_entry.delete(0, tk.END)
                self.api_entry.insert(0, api_key)
                self.log("已從 WebOcrAPI.json 自動載入 API Key")
            else:
                self.log("警告：WebOcrAPI.json 中未找到 API Key 欄位，請手動輸入")
            self.base_url = config.get("base_url") or config.get("url") or "https://api.cloud.llamaindex.ai/api/v2"
        except Exception as e:
            self.log(f"載入 WebOcrAPI.json 失敗：{e}，請手動輸入 API Key")

    def select_dir(self):
        path = filedialog.askdirectory(title="選擇輸出目錄")
        if path:
            self.dir_entry.delete(0, tk.END)
            self.dir_entry.insert(0, path)

    def select_file(self):
        path = filedialog.askopenfilename(
            title="選擇圖片檔案",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.tiff *.pdf")]
        )
        if path:
            self.image_path = path
            self.file_label.config(text=os.path.basename(path), fg='white')
            self.log("已選擇檔案：" + path)

    def log(self, msg):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, msg + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    def process(self):
        api_key = self.api_entry.get().strip()
        if not api_key:
            messagebox.showerror("錯誤", "請先輸入 API Key")
            return
        if not self.image_path or not os.path.exists(self.image_path):
            messagebox.showerror("錯誤", "請先選擇有效的圖片檔案")
            return

        if not (self.var_word.get() or self.var_excel.get() or self.var_html.get() or self.var_md.get()):
            messagebox.showerror("錯誤", "請至少勾選一種輸出格式")
            return

        out_dir = self.dir_entry.get().strip()
        if not out_dir:
            out_dir = os.getcwd()
        if not os.path.exists(out_dir):
            try:
                os.makedirs(out_dir)
            except Exception as e:
                messagebox.showerror("錯誤", f"無法建立輸出目錄：{e}")
                return

        self.process_btn.config(state='disabled')
        self.log("開始處理...")

        try:
            self.log("正在上傳圖片並呼叫 LlamaCloud API v2...")
            response_json = self.call_llamacloud_api_v2(api_key, self.image_path)
            self.log("API 呼叫成功，正在解析 JSON...")

            md_content = self.extract_markdown_content(response_json)
            self.log(f"提取到的內容長度：{len(md_content)} 字元")

            base_name = os.path.splitext(os.path.basename(self.image_path))[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            if self.var_html.get():
                html_filename = f"{base_name}_{timestamp}.html"
                html_path = os.path.join(out_dir, html_filename)
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                self.log(f"原始 HTML 已儲存：{html_path}")

            if self.var_md.get():
                md_filename = f"{base_name}_{timestamp}.md"
                md_path = os.path.join(out_dir, md_filename)
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                self.log(f"MD 文件已儲存：{md_path}")

            # ---- 解析表格（支援 Markdown 與 HTML） ----
            df = None
            try:
                self.log("正在解析表格（先嘗試 Markdown，再嘗試 HTML）...")
                df = self.extract_table_to_dataframe(md_content)
                # 過濾空行（第一欄為空或全為 NaN 的列）
                if df is not None and not df.empty:
                    # 移除第一欄為空字串或 NaN 的列
                    first_col = df.iloc[:, 0]
                    df = df[first_col.notna() & (first_col != '')]
                    # 若仍有全為 NaN 的列則移除
                    df = df.dropna(how='all')
                    self.log(f"過濾空行後，表格尺寸：{df.shape[0]} 列 x {df.shape[1]} 欄")
                else:
                    self.log("警告：解析結果無有效表格")
            except Exception as e:
                self.log(f"警告：表格解析失敗 - {str(e)}")
                df = None

            # ---- 其餘處理（Word 命名、儲存等） ----
            before_text, _ = self._split_md_around_table(md_content)
            word_filename = self._generate_word_filename(before_text, base_name, timestamp)
            word_path = os.path.join(out_dir, word_filename) if self.var_word.get() else ""

            if self.var_excel.get() and df is not None and not df.empty:
                excel_filename = f"{base_name}_{timestamp}.xlsx"
                excel_path = os.path.join(out_dir, excel_filename)
                df.to_excel(excel_path, index=False)
                self.log(f"Excel 已儲存：{excel_path}")
            elif self.var_excel.get():
                self.log("略過 Excel：無有效表格資料")

            if self.var_word.get() and df is not None and not df.empty:
                self.save_as_word(df, word_path, md_content)
                self.log(f"Word 已儲存：{word_path}")
            elif self.var_word.get():
                self.log("略過 Word：無有效表格資料")

            if df is None or df.empty:
                messagebox.showinfo("完成", "處理完成！\n注意：未從圖片中辨識到有效表格，僅儲存了原始內容（HTML/MD）。")
            else:
                messagebox.showinfo("完成", f"轉換完成！\n檔案已儲存至：\n{out_dir}")

        except Exception as e:
            error_msg = f"錯誤：{str(e)}"
            self.log(error_msg)
            messagebox.showerror("處理失敗", error_msg)

        finally:
            self.process_btn.config(state='normal')

    # ====== LlamaCloud API v2 核心函數（最終版） ======
    def call_llamacloud_api_v2(self, api_key, image_path):
        """
        使用 LlamaCloud API v2 進行文件解析。
        上傳端點：POST /api/v2/parse/upload
        輪詢端點：GET /api/v2/parse/{job_id}?expand=markdown
        根據實際回應結構（包含 job 物件與 markdown.pages）進行解析。
        """
        base_url = "https://api.cloud.llamaindex.ai/api/v2"
        headers = {"Authorization": f"Bearer {api_key}"}

        # ---- 步驟 1：上傳檔案 ----
        upload_url = f"{base_url}/parse/upload"
        config = {
            "tier": "cost_effective",   # 可改為 fast / agentic / agentic_plus
            "version": "latest"
        }
        with open(image_path, 'rb') as f:
            files = {"file": (os.path.basename(image_path), f)}
            data = {"configuration": json.dumps(config)}
            response = requests.post(upload_url, headers=headers, files=files, data=data)

        self.log(f"上傳回應碼: {response.status_code}")
        if response.status_code != 200:
            raise Exception(f"上傳失敗 (HTTP {response.status_code})：{response.text}")

        job_json = response.json()
        self.log(f"上傳回應: {json.dumps(job_json, indent=2, ensure_ascii=False)}")

        # 上傳回應直接包含 "id"
        job_id = job_json.get("id")
        if not job_id:
            raise Exception(f"未取得 job_id，回應：{job_json}")

        self.log(f"任務已提交，job_id: {job_id}，開始輪詢...")

        # ---- 步驟 2：輪詢任務狀態 ----
        status_url = f"{base_url}/parse/{job_id}"
        max_attempts = 120
        attempt = 0
        while attempt < max_attempts:
            params = {"expand": "markdown"}
            resp = requests.get(status_url, headers=headers, params=params)
            if resp.status_code != 200:
                raise Exception(f"獲取狀態失敗 (HTTP {resp.status_code})：{resp.text}")

            status_json = resp.json()
            # v2 回應可能將狀態放在 "job" 物件中
            job = status_json.get("job", {})
            status = job.get("status")
            # 若 job 中無狀態，嘗試根層級（相容舊版）
            if status is None:
                status = status_json.get("status")
            self.log(f"當前狀態: {status}")

            status_lower = status.lower() if status else ""

            if status_lower == "completed":
                # 提取 Markdown 內容
                markdown_data = status_json.get("markdown")
                if markdown_data:
                    # 若是物件且包含 pages，合併所有頁面
                    if isinstance(markdown_data, dict) and "pages" in markdown_data:
                        pages = markdown_data.get("pages", [])
                        # 按頁碼排序
                        pages.sort(key=lambda x: x.get("page_number", 0))
                        full_md = "\n".join([p.get("markdown", "") for p in pages if p.get("markdown")])
                        if full_md:
                            return {"markdown": {"content": full_md}}
                    # 若是字串直接返回
                    elif isinstance(markdown_data, str):
                        return {"markdown": {"content": markdown_data}}
                # 若 markdown 欄位不存在，嘗試其他可能
                result = status_json.get("result", {})
                content = result.get("markdown") or result.get("content") or result.get("text")
                if content:
                    return {"markdown": {"content": content}}
                # 最後嘗試根層級的 text
                text = status_json.get("text")
                if text:
                    return {"markdown": {"content": text}}

                self.log(f"完整回應（無 markdown 內容）: {json.dumps(status_json, indent=2, ensure_ascii=False)}")
                raise Exception("完成但未找到 markdown 內容")

            elif status_lower in ["failed", "error"]:
                error_msg = job.get("error_message") or status_json.get("error_message") or "未知錯誤"
                raise Exception(f"解析失敗：{error_msg}")

            elif status_lower in ["pending", "running"]:
                time.sleep(2)
                attempt += 1

            else:
                self.log(f"未知狀態回應: {json.dumps(status_json, indent=2, ensure_ascii=False)}")
                raise Exception(f"未知狀態：{status}")

        raise Exception("輪詢逾時，任務未完成")

    # ====== 輔助方法 ======
    def extract_markdown_content(self, json_data):
        try:
            content = json_data.get("markdown", {}).get("content")
            if content is None:
                content = json_data.get("result", {}).get("markdown")
            if content is None:
                content = json_data.get("result", {}).get("content")
            if content is None:
                raise ValueError("無法從 JSON 中找到 markdown content")
            return content
        except Exception as e:
            raise Exception(f"解析 JSON 失敗：{e}\n原始回應：{json_data}")

    def extract_table_to_dataframe(self, text):
        """
        從 text 中提取表格，優先解析 Markdown 表格，若失敗則嘗試 HTML 表格。
        """
        # 1) 嘗試 Markdown 表格（| 分隔）
        md_table = self._extract_markdown_table(text)
        if md_table:
            try:
                df = pd.read_markdown(io.StringIO(md_table))
                return df
            except (ImportError, AttributeError):
                return self._parse_markdown_table(md_table)
            except Exception:
                return self._parse_markdown_table(md_table)

        # 2) 嘗試 HTML 表格（使用 pandas.read_html）
        # 搜尋 <table> 標籤
        if '<table' in text.lower():
            try:
                tables = pd.read_html(io.StringIO(text))
                if tables and len(tables) > 0:
                    df = tables[0]  # 取第一個表格
                    # 若標題列重複（有些 HTML 表格會把 th 也當成資料），可手動檢查
                    return df
            except Exception as e:
                # 若 pandas 解析失敗，使用正則或 BeautifulSoup 備用（此處略）
                pass

        raise ValueError("未找到任何 Markdown 或 HTML 表格")

    def _extract_markdown_table(self, text):
        lines = text.splitlines()
        table_lines = []
        in_table = False
        for line in lines:
            if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
                if not in_table:
                    in_table = True
                table_lines.append(line)
            else:
                if in_table:
                    if len(table_lines) >= 3:
                        break
                    else:
                        table_lines = []
                        in_table = False
        if len(table_lines) >= 3:
            return '\n'.join(table_lines)
        return None

    def _parse_markdown_table(self, md_text):
        lines = md_text.splitlines()
        data_lines = []
        for line in lines:
            if re.search(r'^\s*\|[\s\-:]+?\|', line):
                continue
            if '|' in line:
                stripped = line.strip('|')
                cells = [cell.strip() for cell in stripped.split('|')]
                data_lines.append(cells)
        if not data_lines:
            raise ValueError("沒有有效的資料行")
        header = data_lines[0]
        data = data_lines[1:]
        max_cols = max(len(row) for row in data_lines)
        for row in data:
            while len(row) < max_cols:
                row.append('')
        while len(header) < max_cols:
            header.append('')
        return pd.DataFrame(data, columns=header[:max_cols])

    def _split_md_around_table(self, md_text):
        """
        傳回 (表格前文字, 表格後文字)，完整排除表格本身（支援 HTML 與 Markdown 表格）。
        修正：優先使用正則匹配完整 HTML 表格，避免殘留 HTML 標籤。
        """
        # 1. 優先處理 HTML 表格（<table> ... </table>）
        html_pattern = re.compile(r'<table[^>]*>.*?</table>', re.IGNORECASE | re.DOTALL)
        html_match = html_pattern.search(md_text)
        if html_match:
            start, end = html_match.span()
            before = md_text[:start].strip()
            after = md_text[end:].strip()
            return before, after

        # 2. 處理 Markdown 表格（連續的 | 分隔行）
        lines = md_text.splitlines()
        table_start = None
        table_end = None
        in_table = False

        for i, line in enumerate(lines):
            if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
                if not in_table:
                    in_table = True
                    table_start = i
            else:
                if in_table:
                    table_end = i
                    break

        if table_start is not None:
            if table_end is None:
                table_end = len(lines)
            before = '\n'.join(lines[:table_start]).strip()
            after = '\n'.join(lines[table_end:]).strip()
            return before, after

        # 沒有任何表格
        return md_text, ""

    def _generate_word_filename(self, before_text, fallback_base, timestamp):
        date_str = None
        warehouse_str = None

        date_match = re.search(r'(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', before_text)
        if date_match:
            y, m, d = date_match.groups()
            date_str = f"{y}年{m}月{d}日"

        wh_match = re.search(r'\((\d{1,3})\s*庫\)', before_text)
        if wh_match:
            warehouse_str = f"({wh_match.group(1)}庫)"

        if date_str and warehouse_str:
            return f"{date_str}{warehouse_str}.docx"
        else:
            return f"{fallback_base}_{timestamp}.docx"

    def _merge_title_info(self, lines):
        if not lines:
            return lines

        new_lines = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if '庫)' in line and (line.startswith('#') or '(' in line):
                combined_parts = [line]
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if any(next_line.startswith(prefix) for prefix in [
                        '**日期：**', '日期：',
                        '**填單者：**', '填單者：',
                        '**單位：**', '單位：'
                    ]):
                        combined_parts.append(next_line)
                        j += 1
                    else:
                        break
                merged_line = '     '.join(combined_parts)
                new_lines.append(merged_line)
                i = j
            else:
                new_lines.append(line)
                i += 1
        return new_lines

    def save_as_word(self, df, filename, md_content):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        # 取得表格前後文字（利用 HTML 或 Markdown 分隔）
        before_text, after_text = self._split_md_around_table(md_content)

        if before_text:
            before_lines = [line.strip() for line in before_text.split('\n') if line.strip()]
            merged_lines = self._merge_title_info(before_lines)
            for line in merged_lines:
                doc.add_paragraph(line)

        if df.shape[1] > 0:
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Table Grid'
            for j, col_name in enumerate(df.columns):
                table.cell(0, j).text = str(col_name)
            for i, row in df.iterrows():
                for j, value in enumerate(row):
                    cell_text = "" if pd.isna(value) else str(value)
                    table.cell(i + 1, j).text = cell_text

        if after_text:
            doc.add_paragraph()
            for line in after_text.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

        doc.save(filename)

    # 以下為保留的輔助函式（未使用但保留以防萬一）
    def _clean_html_text(self, raw_text):
        decoded = html.unescape(raw_text)
        plain = re.sub(r'<[^>]+>', '', decoded)
        lines = [line.strip() for line in plain.splitlines() if line.strip()]
        return '\n'.join(lines)

    def extract_text_before_first_table(self, html_str):
        match = re.search(r'<table.*?>', html_str, re.IGNORECASE)
        if not match:
            return ""
        before = html_str[:match.start()]
        return self._clean_html_text(before)

    def extract_text_after_first_table(self, html_str):
        match = re.search(r'<table.*?>.*?</table>', html_str, re.DOTALL | re.IGNORECASE)
        if not match:
            return ""
        after = html_str[match.end():]
        return self._clean_html_text(after)


if __name__ == "__main__":
    root = tk.Tk()
    app = LlamaCloudOCRApp(root)
    root.mainloop()