import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Checkbutton, IntVar
import requests
import pandas as pd
from docx import Document
from docx.shared import Cm
import os
import re
import io
import html
import json
from datetime import datetime
from PIL import Image

class OCRSpaceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("OCR.Space 表格提取工具")
        self.root.geometry("650x720")
        self.root.configure(bg='black')

        fg_color = 'white'
        bg_color = 'black'
        entry_bg = '#333333'
        btn_bg = '#555555'

        title = tk.Label(root, text="OCR.Space 表格提取工具", font=('Arial', 16, 'bold'),
                         fg=fg_color, bg=bg_color)
        title.pack(pady=10)

        # API Key 狀態標籤 (先建立，但不在此時載入 key)
        self.api_key = None
        self.api_status_label = tk.Label(root, text="API Key：尚未載入", fg='orange', bg=bg_color,
                                         font=('Arial', 10))
        self.api_status_label.pack(pady=2)

        # 輸出目錄選擇
        frame_out = tk.Frame(root, bg=bg_color)
        frame_out.pack(pady=5, fill='x', padx=20)
        tk.Label(frame_out, text="輸出目錄：", fg=fg_color, bg=bg_color,
                 font=('Arial', 10)).pack(side='left')
        self.dir_entry = tk.Entry(frame_out, width=40, bg=entry_bg, fg=fg_color)
        self.dir_entry.pack(side='left', padx=5)
        self.dir_entry.insert(0, r"E:\DiskCUse\HFDownloads")
        tk.Button(frame_out, text="瀏覽", command=self.select_dir,
                  bg=btn_bg, fg=fg_color).pack(side='left', padx=5)

        # 圖片選擇
        frame_file = tk.Frame(root, bg=bg_color)
        frame_file.pack(pady=10, fill='x', padx=20)
        self.file_label = tk.Label(frame_file, text="未選擇圖片", fg='gray', bg=bg_color,
                                   width=40, anchor='w')
        self.file_label.pack(side='left')
        tk.Button(frame_file, text="選擇圖片", command=self.select_file,
                  bg=btn_bg, fg=fg_color, padx=10).pack(side='right')

        # 輸出格式勾選框
        frame_options = tk.LabelFrame(root, text="儲存格式選項", bg=bg_color, fg=fg_color,
                                      font=('Arial', 10))
        frame_options.pack(pady=10, fill='x', padx=20)

        self.var_word = IntVar(value=1)
        self.var_excel = IntVar(value=0)
        self.var_text = IntVar(value=1)
        self.var_md = IntVar(value=1)

        chk_word = Checkbutton(frame_options, text="儲存 WORD 檔", variable=self.var_word,
                               bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_word.pack(side='left', padx=10)

        chk_excel = Checkbutton(frame_options, text="儲存 EXCEL 檔", variable=self.var_excel,
                                bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_excel.pack(side='left', padx=10)

        chk_text = Checkbutton(frame_options, text="儲存原始 OCR 文字檔 (TXT)", variable=self.var_text,
                               bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_text.pack(side='left', padx=10)

        chk_md = Checkbutton(frame_options, text="儲存 MD 文件", variable=self.var_md,
                             bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_md.pack(side='left', padx=10)

        # 表格辨識選項 (isTable)
        self.var_istable = IntVar(value=1)
        chk_istable = Checkbutton(frame_options, text="表格/收據辨識 (isTable)", variable=self.var_istable,
                                  bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_istable.pack(side='left', padx=10)

        # 開始按鈕
        self.process_btn = tk.Button(root, text="🚀 開始辨識並轉換", command=self.process,
                                     bg='#0066cc', fg=fg_color, font=('Arial', 12, 'bold'),
                                     padx=20, pady=5)
        self.process_btn.pack(pady=15)

        # 日誌區域
        self.log_area = scrolledtext.ScrolledText(root, height=18, bg='#1a1a1a',
                                                  fg='#00ff00', font=('Consolas', 10))
        self.log_area.pack(fill='both', expand=True, padx=20, pady=10)
        self.log_area.insert(tk.END, "就緒，從 WebOcrAPI.json 載入 API Key。\n")
        self.log_area.config(state='disabled')

        # 在這裡才載入 API Key，確保 log_area 已可用
        self.load_api_key()

        self.image_path = None

    def load_api_key(self):
        """從 WebOcrAPI.json 載入 ocr.space 模型的 API Key"""
        json_path = "WebOcrAPI.json"
        if not os.path.exists(json_path):
            self.log(f"錯誤：找不到設定檔 {json_path}")
            self.api_status_label.config(text="API Key：設定檔遺失", fg='red')
            return
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                configs = json.load(f)
        except Exception as e:
            self.log(f"錯誤：讀取設定檔失敗 - {e}")
            self.api_status_label.config(text="API Key：讀取失敗", fg='red')
            return

        for cfg in configs:
            if cfg.get("ModelName") == "ocr.space":
                key = cfg.get("api_key")
                if key:
                    self.api_key = key
                    self.api_status_label.config(text="API Key：已載入", fg='#00ff00')
                    self.log("API Key 已從 WebOcrAPI.json 載入")
                    return
                else:
                    self.api_status_label.config(text="API Key：金鑰為空", fg='red')
                    self.log("錯誤：設定檔中 ocr.space 的 api_key 為空")
                    return
        self.api_status_label.config(text="API Key：找不到模型", fg='red')
        self.log("錯誤：設定檔中找不到模型 ocr.space")

    def select_dir(self):
        path = filedialog.askdirectory(title="選擇輸出目錄")
        if path:
            self.dir_entry.delete(0, tk.END)
            self.dir_entry.insert(0, path)

    def select_file(self):
        path = filedialog.askopenfilename(
            title="選擇圖片檔案",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.tiff")]
        )
        if path:
            self.image_path = path
            self.file_label.config(text=os.path.basename(path), fg='white')
            self.log("已選擇檔案：" + path)

    def log(self, msg):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, msg + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    def process(self):
        if not self.api_key:
            messagebox.showerror("錯誤", "API Key 未正確載入，請檢查 WebOcrAPI.json")
            return
        if not self.image_path or not os.path.exists(self.image_path):
            messagebox.showerror("錯誤", "請先選擇有效的圖片檔案")
            return

        if not (self.var_word.get() or self.var_excel.get() or self.var_text.get() or self.var_md.get()):
            messagebox.showerror("錯誤", "請至少勾選一種輸出格式")
            return

        out_dir = self.dir_entry.get().strip()
        if not out_dir:
            out_dir = os.getcwd()
        if not os.path.exists(out_dir):
            try:
                os.makedirs(out_dir)
            except Exception as e:
                messagebox.showerror("錯誤", f"無法建立輸出目錄：{e}")
                return

        self.process_btn.config(state='disabled')
        self.log("開始處理...")

        try:
            self.log("正在上傳圖片並呼叫 OCR.Space API...")
            isTable = bool(self.var_istable.get())
            ocr_text = self.call_ocrspace_api(self.api_key, self.image_path,
                                              language='cht',
                                              isOverlayRequired=True,
                                              isTable=isTable,
                                              OCREngine=3)
            self.log(f"OCR 成功，提取文字長度：{len(ocr_text)} 字元")

            # ---- 決定基礎檔名：優先使用日期+庫別 ----
            date_str, warehouse_str = self.extract_date_and_warehouse(ocr_text)
            if date_str or warehouse_str:
                base_filename = f"{date_str}{warehouse_str}"
                self.log(f"從文字中提取檔名：{base_filename}")
            else:
                base_name = os.path.splitext(os.path.basename(self.image_path))[0]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_filename = f"{base_name}_{timestamp}"
                self.log("未偵測到日期或庫別，使用預設檔名")

            # 儲存原始文字檔 (TXT)
            if self.var_text.get():
                txt_filename = f"{base_filename}.txt"
                txt_path = os.path.join(out_dir, txt_filename)
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(ocr_text)
                self.log(f"原始文字檔已儲存：{txt_path}")

            # 儲存 MD 文件
            if self.var_md.get():
                md_filename = f"{base_filename}.md"
                md_path = os.path.join(out_dir, md_filename)
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(ocr_text)
                self.log(f"MD 文件已儲存：{md_path}")

            # ---- 表格提取與修正（支援 PNG 轉換重試） ----
            MAX_RETRIES = 1
            attempt = 0
            df = None
            before_text = ""
            after_text = ""
            final_ocr_text = ocr_text

            while attempt <= MAX_RETRIES:
                try:
                    self.log(f"第 {attempt+1} 次解析表格...")
                    before_text, table_text, after_text = self.split_text_around_first_table(final_ocr_text)
                    if not table_text:
                        self.log("未偵測到表格區塊")
                        break

                    df = self.extract_table_to_dataframe_from_block(table_text)
                    self.log(f"表格解析成功，尺寸：{df.shape[0]} 列 x {df.shape[1]} 欄")

                    if df is not None:
                        df = self.fix_table_columns(df)
                        if df.shape[1] == 8:
                            self.log("表格欄數正確（8 欄）")
                            break
                        else:
                            self.log(f"表格欄數不正確 ({df.shape[1]} 欄)，需要重新辨識")
                    else:
                        self.log("表格解析結果為空")
                except Exception as e:
                    self.log(f"表格處理失敗：{str(e)}")

                attempt += 1
                if attempt <= MAX_RETRIES:
                    self.log("正在將圖片轉換為 PNG 格式並重新呼叫 API...")
                    try:
                        png_path = self.convert_to_png(self.image_path)
                        final_ocr_text = self.call_ocrspace_api(
                            self.api_key,
                            png_path,
                            language='cht',
                            isOverlayRequired=True,
                            isTable=isTable,
                            OCREngine=3
                        )
                        self.log(f"重新辨識成功，文字長度：{len(final_ocr_text)} 字元")
                        if png_path != self.image_path:
                            os.remove(png_path)
                    except Exception as e:
                        self.log(f"PNG 轉換或重新辨識失敗：{e}")
                        break
            else:
                self.log("已達最大重試次數，將使用最後一次的表格結果（可能不完整）")

            # 儲存 Excel
            if self.var_excel.get() and df is not None:
                excel_filename = f"{base_filename}.xlsx"
                excel_path = os.path.join(out_dir, excel_filename)
                df.to_excel(excel_path, index=False)
                self.log(f"Excel 已儲存：{excel_path}")
            elif self.var_excel.get() and df is None:
                self.log("略過 Excel：未找到表格資料")

            # 儲存 Word
            if self.var_word.get():
                word_filename = f"{base_filename}.docx"
                word_path = os.path.join(out_dir, word_filename)
                self.save_as_word(df, word_path, before_text, after_text)
                self.log(f"Word 已儲存：{word_path}")

            if df is None and before_text.strip() == "" and after_text.strip() == "":
                messagebox.showinfo("完成", "處理完成，但未辨識到表格或文字。")
            else:
                messagebox.showinfo("完成", f"轉換完成！\n檔案已儲存至：\n{out_dir}")

        except Exception as e:
            error_msg = f"錯誤：{str(e)}"
            self.log(error_msg)
            messagebox.showerror("處理失敗", error_msg)

        finally:
            self.process_btn.config(state='normal')

    def call_ocrspace_api(self, api_key, image_path, language='cht',
                          isOverlayRequired=True, isTable=True, OCREngine=3):
        url = "https://api.ocr.space/parse/image"
        headers = {"apikey": api_key}
        with open(image_path, 'rb') as f:
            files = {"file": (os.path.basename(image_path), f)}
            data = {
                "language": language,
                "isOverlayRequired": isOverlayRequired,
                "isTable": isTable,
                "OCREngine": OCREngine,
            }
            response = requests.post(url, headers=headers, files=files, data=data)

        if response.status_code != 200:
            raise Exception(f"API 請求失敗 (HTTP {response.status_code})：{response.text}")

        result = response.json()
        if result.get("IsErroredOnProcessing") or result.get("OCRExitCode") != 1:
            error_details = ""
            if result.get("ParsedResults"):
                parsed = result["ParsedResults"][0]
                error_details = parsed.get("ErrorDetails", "") or parsed.get("ErrorMessage", "")
            raise Exception(f"OCR 處理失敗：{error_details or result}")

        parsed_results = result.get("ParsedResults")
        if not parsed_results or len(parsed_results) == 0:
            raise Exception("未返回任何 OCR 結果")

        parsed_text = parsed_results[0].get("ParsedText", "")
        if not parsed_text:
            raise Exception("OCR 結果為空")
        return parsed_text

    def extract_date_and_warehouse(self, text):
        # 支援 (4-2庫) 這種格式
        wh_pattern = r'(\(\d+(-\d+)?庫\))'
        wh_match = re.search(wh_pattern, text)
        warehouse = wh_match.group(0) if wh_match else ""

        date_pattern = r'((\d{2,4})年\s*\d{1,2}月\s*\d{1,2}日)'
        date_match = re.search(date_pattern, text)
        date = date_match.group(0) if date_match else ""
        if date:
            date = re.sub(r'\s+', '', date)
        return date, warehouse

    def split_text_around_first_table(self, text):
        # 嘗試 HTML 表格
        html_pattern = re.compile(r'(<table.*?>.*?</table>)', re.DOTALL | re.IGNORECASE)
        html_match = html_pattern.search(text)
        if html_match:
            table_block = html_match.group(1)
            before = text[:html_match.start()]
            after = text[html_match.end():]
            before_clean = self._html_to_plaintext(before)
            after_clean = self._html_to_plaintext(after)
            return before_clean, table_block, after_clean

        # 嘗試 Markdown 表格
        md_pattern = r'(^|\n)((?:[^\n]*\|[^\n]*\n)+)'
        md_match = re.search(md_pattern, text, re.MULTILINE)
        if md_match:
            table_block = md_match.group(2).strip()
            before = text[:md_match.start()]
            after = text[md_match.end():]
            before_clean = before.strip()
            after_clean = after.strip()
            return before_clean, table_block, after_clean

        return text, "", ""

    def _html_to_plaintext(self, html_str):
        decoded = html.unescape(html_str)
        plain = re.sub(r'<br\s*/?>', '\n', decoded, flags=re.IGNORECASE)
        plain = re.sub(r'<[^>]+>', '', plain)
        lines = [line.strip() for line in plain.splitlines() if line.strip()]
        return '\n'.join(lines)

    def extract_table_to_dataframe_from_block(self, block_text):
        if block_text.strip().startswith('<table'):
            dfs = pd.read_html(io.StringIO(block_text))
            if dfs:
                return dfs[0]
            else:
                raise ValueError("HTML 表格解析失敗")
        else:
            try:
                df = pd.read_markdown(io.StringIO(block_text))
                return df
            except (ImportError, AttributeError):
                return self._parse_markdown_table(block_text)
            except Exception:
                return self._parse_markdown_table(block_text)

    def _parse_markdown_table(self, md_text):
        lines = md_text.splitlines()
        data_lines = []
        for line in lines:
            if re.search(r'^\s*\|[\s\-:]+?\|', line):
                continue
            if '|' in line:
                stripped = line.strip('|')
                cells = [cell.strip() for cell in stripped.split('|')]
                data_lines.append(cells)
        if not data_lines:
            raise ValueError("沒有有效的資料行")
        header = data_lines[0]
        data = data_lines[1:]
        max_cols = max(len(row) for row in data_lines)
        for row in data:
            while len(row) < max_cols:
                row.append('')
        while len(header) < max_cols:
            header.append('')
        df = pd.DataFrame(data, columns=header[:max_cols])
        return df

    def fix_table_columns(self, df):
        if df is None or df.empty:
            return df
        df.columns = [str(col).strip() for col in df.columns]
        if len(df.columns) == 8:
            if df.columns[0] == '' or '品項' not in df.columns[0]:
                new_cols = list(df.columns)
                new_cols[0] = '品項'
                df.columns = new_cols
                self.log("已修正第一欄名稱為「品項」")
            return df
        return df

    def convert_to_png(self, image_path):
        img = Image.open(image_path)
        if img.format == 'PNG':
            return image_path
        base, _ = os.path.splitext(image_path)
        png_path = base + "_temp_ocr.png"
        img = img.convert('RGB')
        img.save(png_path, 'PNG')
        self.log(f"圖片已轉換為 PNG：{png_path}")
        return png_path

    def save_as_word(self, df, filename, before_text, after_text):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        date_pattern = re.compile(r'(\d{2,4}年\s*\d{1,2}月\s*\d{1,2}日)')
        # 支援 (4-2庫) 等含連字號的庫別格式
        warehouse_pattern = re.compile(r'(\(\d+(-\d+)?庫\).*進、銷貨庫存表)')
        filler_pattern = re.compile(r'填單者:\s*(.*)')
        unit_pattern = re.compile(r'單位:\s*(.*)')

        title = ""
        date_str = ""
        filler = ""
        unit = ""

        lines = before_text.splitlines()
        for line in lines:
            w_match = warehouse_pattern.search(line)
            if w_match:
                title = w_match.group(0).strip()
                continue
            d_match = date_pattern.search(line)
            if d_match:
                date_str = d_match.group(0).replace(' ', '')
            f_match = filler_pattern.search(line)
            if f_match:
                filler = f_match.group(0).strip()
            u_match = unit_pattern.search(line)
            if u_match:
                unit = u_match.group(0).strip()

        if title and date_str and filler and unit:
            combined_text = f"{title}\t{date_str}\t{filler}\t{unit}"
            header_paragraph = doc.add_paragraph(combined_text)
            tab_stops = header_paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(6))
            tab_stops.add_tab_stop(Cm(10))
            tab_stops.add_tab_stop(Cm(14))
            remaining_before_lines = [
                l for l in lines
                if not (warehouse_pattern.search(l) or date_pattern.search(l) or
                        filler_pattern.search(l) or unit_pattern.search(l))
            ]
        else:
            remaining_before_lines = lines

        for line in remaining_before_lines:
            if line.strip():
                doc.add_paragraph(line.strip())

        if df is not None and not df.empty:
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Table Grid'
            for j, col_name in enumerate(df.columns):
                table.cell(0, j).text = str(col_name)
            for i, row in df.iterrows():
                for j, value in enumerate(row):
                    if pd.isna(value):
                        cell_text = ""
                    else:
                        cell_text = str(value)
                    table.cell(i + 1, j).text = cell_text

        if after_text.strip():
            for line in after_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

        doc.save(filename)


if __name__ == "__main__":
    root = tk.Tk()
    app = OCRSpaceApp(root)
    root.mainloop()