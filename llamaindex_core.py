#!/usr/bin/env python3
"""
LlamaIndex OCR Core Module
Contains the core OCR processing logic for LlamaCloud API, separated from UI.
"""

import os
import re
import json
import time
from pathlib import Path
from datetime import datetime
from typing import Dict, Tuple, Optional, Any

import requests
import pandas as pd
from docx import Document
from docx.shared import Cm
import io


def load_api_key(json_path: str = "WebOcrAPI.json") -> tuple[str, str]:
    """Load LlamaCloud API key and base URL from WebOcrAPI.json.
    If the file doesn't exist, creates a default one with a placeholder key.

    Args:
        json_path: Path to the JSON configuration file.

    Returns:
        Tuple of (api_key, base_url).

    Raises:
        ValueError: If the API key is missing or malformed.
    """
    if not os.path.exists(json_path):
        # Create a default config file
        default_config = {
            "url": "https://api.cloud.llamaindex.ai",
            "base_url": "https://api.cloud.llamaindex.ai/api/v2",
            "api_key": "llx-kp6L7gAeSDVvtfS4ZXXwkFeQNG82KfcoSxpoC1Wqga1cr5Xa"  # placeholder
        }
        try:
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(default_config, f, indent=2, ensure_ascii=False)
        except Exception as e:
            raise ValueError(f"Failed to create default config file: {e}")

    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
    except json.JSONDecodeError as e:
        raise ValueError(f"Invalid JSON in {json_path}: {e}")

    api_key = config.get("api_key") or config.get("LLAMA_CLOUD_API_KEY") or config.get("API_KEY")
    if not api_key:
        raise ValueError("API key for LlamaCloud is empty in WebOcrAPI.json")

    base_url = config.get("base_url") or config.get("url") or "https://api.cloud.llamaindex.ai/api/v2"
    return api_key, base_url


def call_llamacloud_api(api_key: str, image_path: str, base_url: str) -> dict:
    """Call LlamaCloud API v2 for document parsing.

    Args:
        api_key: LlamaCloud API key.
        image_path: Path to the image file to process.
        base_url: Base URL for the API (e.g., "https://api.cloud.llamaindex.ai/api/v2").

    Returns:
        The JSON response from the API.

    Raises:
        Exception: If the API returns an error or the request fails.
    """
    upload_url = f"{base_url}/parse/upload"
    headers = {"Authorization": f"Bearer {api_key}"}
    
    # Configuration for the processing tier
    config = {
        "tier": "cost_effective",  # can be fast, agentic, agentic_plus
        "version": "latest"
    }
    
    with open(image_path, 'rb') as f:
        files = {"file": (os.path.basename(image_path), f)}
        data = {"configuration": json.dumps(config)}
        response = requests.post(upload_url, headers=headers, files=files, data=data, timeout=120)

    if response.status_code != 200:
        raise Exception(f"Upload failed (HTTP {response.status_code}): {response.text}")

    job_json = response.json()
    job_id = job_json.get("id")
    if not job_id:
        raise Exception(f"Failed to get job ID from response: {job_json}")

    # Poll for completion
    status_url = f"{base_url}/parse/{job_id}"
    params = {"expand": "markdown"}
    start_time = time.time()
    while time.time() - start_time < 120:  # 2 minute timeout
        resp = requests.get(status_url, headers=headers, params=params, timeout=30)
        if resp.status_code != 200:
            raise Exception(f"Failed to get status (HTTP {resp.status_code}): {resp.text}")
        
        status_json = resp.json()
        # The response may have a 'job' object or direct fields
        job = status_json.get("job", {})
        status = job.get("status") or status_json.get("status")
        
        if status == "completed":
            # Extract markdown content
            markdown_data = status_json.get("markdown") or job.get("markdown")
            if markdown_data:
                if isinstance(markdown_data, dict) and "pages" in markdown_data:
                    pages = markdown_data.get("pages", [])
                    pages.sort(key=lambda x: x.get("page_number", 0))
                    full_md = "\n".join([p.get("markdown", "") for p in pages if p.get("markdown")])
                    if full_md:
                        return {"markdown": {"content": full_md}}
                elif isinstance(markdown_data, str):
                    return {"markdown": {"content": markdown_data}}
            
            # Fallback: try to get content from result or text fields
            result = status_json.get("result", {})
            content = result.get("markdown") or result.get("content") or result.get("text")
            if content:
                return {"markdown": {"content": content}}
            
            text = status_json.get("text") or job.get("text")
            if text:
                return {"markdown": {"content": text}}
                
            raise Exception("Completed but no markdown content found")
        
        elif status in ("failed", "error"):
            error_msg = job.get("error_message") or status_json.get("error_message") or "Unknown error"
            raise Exception(f"Processing failed: {error_msg}")
        
        # Still processing, wait a bit
        time.sleep(2)
    
    raise Exception("Polling timeout - processing did not complete within 2 minutes")


def extract_markdown_content(json_data: dict) -> str:
    """Extract markdown content from LlamaCloud API response.

    Args:
        json_data: The JSON response from LlamaCloud API.

    Returns:
        The markdown/content string.

    Raises:
        ValueError: If markdown content cannot be found.
    """
    try:
        content = json_data.get("markdown", {}).get("content")
        if content is None:
            # Try alternative paths
            content = json_data.get("result", {}).get("markdown")
        if content is None:
            content = json_data.get("result", {}).get("content")
        if content is None:
            raise ValueError("Unable to find markdown content in JSON")
        return content
    except Exception as e:
        raise ValueError(f"Failed to parse JSON: {e}")


def extract_table_to_dataframe(markdown_text: str) -> pd.DataFrame | None:
    """Extract the first table from markdown text and convert to DataFrame.

    Tries to parse as a markdown table first, then falls back to HTML if markdown fails.

    Args:
        markdown_text: The markdown text that may contain a table.

    Returns:
        pandas DataFrame if a table is found, None otherwise.
    """
    # Try to find a markdown table (lines with |)
    lines = markdown_text.splitlines()
    table_lines = []
    in_table = False
    
    for line in lines:
        # Check if line looks like a markdown table row
        if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
            if not in_table:
                in_table = True
            table_lines.append(line)
        else:
            if in_table:
                # End of table
                if len(table_lines) >= 3:  # Need at least header, separator, and one data row
                    break
                else:
                    # Reset if we didn't get a proper table
                    table_lines = []
                    in_table = False
    
    if len(table_lines) >= 3:
        table_text = "\n".join(table_lines)
        try:
            # Try using pandas' read_markdown if available
            return pd.read_markdown(io.StringIO(table_text))
        except Exception:
            # Fallback to manual parsing
            return _parse_markdown_table(table_text)
    
    # If no markdown table found, try to find an HTML table
    html_pattern = re.compile(r'<table[^>]*>.*?</table>', re.IGNORECASE | re.DOTALL)
    html_match = html_pattern.search(markdown_text)
    if html_match:
        html_table = html_match.group(0)
        try:
            dfs = pd.read_html(io.StringIO(html_table))
            if dfs:
                return dfs[0]
        except Exception:
            pass
    
    return None


def _parse_markdown_table(md_text: str) -> pd.DataFrame:
    """Parse a markdown table string into a DataFrame (fallback implementation).

    Args:
        md_text: Markdown table string.

    Returns:
        pandas DataFrame.
    """
    lines = md_text.splitlines()
    data_lines = []
    for line in lines:
        # Skip separator lines (e.g., |---|---|)
        if re.search(r'^\s*\|[\s\-:]+?\|', line):
            continue
        if '|' in line:
            # Remove leading/trailing pipes and split
            stripped = line.strip('|')
            cells = [cell.strip() for cell in stripped.split('|')]
            data_lines.append(cells)
    
    if not data_lines:
        raise ValueError("No valid data rows found in markdown table")
    
    header = data_lines[0]
    data = data_lines[1:]
    
    # Determine maximum number of columns
    max_cols = max(len(row) for row in data_lines)
    
    # Pad rows that are shorter than max_cols
    for row in data:
        while len(row) < max_cols:
            row.append('')
    
    # Pad header if needed
    while len(header) < max_cols:
        header.append('')
    
    return pd.DataFrame(data, columns=header[:max_cols])


def split_text_around_table(text: str) -> tuple[str, str]:
    """Split text into parts before and after the first table (markdown or HTML).

    Args:
        text: The text that may contain a table.

    Returns:
        Tuple (before_text, after_text). If no table is found, returns (text, "").
    """
    # Try to find an HTML table first
    html_pattern = re.compile(r'<table[^>]*>.*?</table>', re.IGNORECASE | re.DOTALL)
    html_match = html_pattern.search(text)
    if html_match:
        start, end = html_match.span()
        before = text[:start]
        after = text[end:]
        return before.strip(), after.strip()
    
    # Fallback to markdown table detection
    lines = text.splitlines()
    table_start = None
    table_end = None
    in_table = False
    
    for i, line in enumerate(lines):
        if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
            if not in_table:
                in_table = True
                table_start = i
        else:
            if in_table:
                table_end = i
                break
    
    if table_start is not None:
        if table_end is None:
            table_end = len(lines)
        before_lines = lines[:table_start]
        after_lines = lines[table_end:]
        before_text = "\n".join(before_lines).strip()
        after_text = "\n".join(after_lines).strip()
        return before_text, after_text
    
    return text, ""


def _generate_word_filename(before_text: str, fallback_base: str, timestamp: str) -> str:
    """Generate a Word filename based on extracted date and warehouse info.

    Args:
        before_text: Text before the table (may contain date and warehouse).
        fallback_base: Base name to use if extraction fails.
        timestamp: Timestamp string for fallback.

    Returns:
        Suggested filename (without directory).
    """
    date_str = None
    warehouse_str = None
    
    # Date: support Chinese year/month/day with optional spaces
    date_match = re.search(r'(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', before_text)
    if date_match:
        y, m, d = date_match.groups()
        date_str = f"{y}年{m}月{d}日"
    
    # Warehouse: support patterns like (6庫) or (4-2庫)
    wh_match = re.search(r'\((\d{1,3}(?:-\d{1,2})?)\s*庫\)', before_text)
    if wh_match:
        warehouse_str = f"({wh_match.group(1)}庫)"
    
    if date_str and warehouse_str:
        return f"{date_str}{warehouse_str}.docx"
    else:
        return f"{fallback_base}_{timestamp}.docx"


def _merge_title_info(lines: list[str]) -> list[str]:
    """Merge header information lines (warehouse, date, filler, unit) into a single line.

    Args:
        lines: List of lines before the table.

    Returns:
        List of lines with header info merged.
    """
    if not lines:
        return lines
    
    new_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if '庫)' in line and (line.startswith('#') or '(' in line):
            combined = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if any(next_line.startswith(prefix) for prefix in [
                    '**日期：**', '日期：',
                    '**填單者：**', '填單者：',
                    '**單位：**', '單位：'
                ]):
                    combined.append(next_line)
                    j += 1
                else:
                    break
            merged_line = '     '.join(combined)
            new_lines.append(merged_line)
            i = j
        else:
            new_lines.append(line)
            i += 1
    return new_lines


def save_as_word(df: pd.DataFrame, image_path: str, output_dir: str, before_text: str) -> str:
    """Save DataFrame as a Word document with header info extracted from text.

    Args:
        df: DataFrame containing the table data.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.
        before_text: Text before the table (used for header info).

    Returns:
        Full path to the saved Word file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    word_filename = f"{base_name}_{timestamp}.docx"
    word_path = os.path.join(output_dir, word_filename)
    
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    
    if before_text:
        before_lines = [line.strip() for line in before_text.split('\n') if line.strip()]
        merged_lines = _merge_title_info(before_lines)
        for line in merged_lines:
            doc.add_paragraph(line)
    
    if df.shape[1] > 0:
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'
        for j, col_name in enumerate(df.columns):
            table.cell(0, j).text = str(col_name)
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                cell_text = "" if pd.isna(value) else str(value)
                table.cell(i + 1, j).text = cell_text
    
    doc.save(word_path)
    return word_path


def save_as_excel(df: pd.DataFrame, image_path: str, output_dir: str) -> str:
    """Save DataFrame as an Excel file.

    Args:
        df: DataFrame containing the table data.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved Excel file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    excel_filename = f"{base_name}_{timestamp}.xlsx"
    excel_path = os.path.join(output_dir, excel_filename)
    df.to_excel(excel_path, index=False)
    return excel_path


def save_as_md(content: str, image_path: str, output_dir: str) -> str:
    """Save markdown content to a .md file.

    Args:
        content: The markdown content to save.
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved markdown file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    md_filename = f"{base_name}_{timestamp}.md"
    md_path = os.path.join(output_dir, md_filename)
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# OCR Results - {os.path.basename(image_path)}\n\n")
        f.write(f"**Processing Time**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write(content)
    return md_path


def save_as_html(content: str, image_path: str, output_dir: str) -> str:
    """Save content to an .html file.

    Args:
        content: The content to save (typically markdown that contains HTML).
        image_path: Path to the original image (used for naming).
        output_dir: Directory where the file should be saved.

    Returns:
        Full path to the saved HTML file.
    """
    base_name = os.path.splitext(os.path.basename(image_path))[0]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    html_filename = f"{base_name}_{timestamp}.html"
    html_path = os.path.join(output_dir, html_filename)
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(f"<!DOCTYPE html>\n<html>\n<head>\n<title>OCR Results - {os.path.basename(image_path)}</title>\n</head>\n<body>\n")
        f.write(f"<h1>OCR Results - {os.path.basename(image_path)}</h1>")
        f.write(f"<p><strong>Processing Time</span>: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")
        f.write(f"<div>{content}</div>")
        f.write("\n</body>\n</html>")
    return html_path


def process_image(api_key: str, image_path: str, output_dir: str,
                  save_word: bool = True, save_excel: bool = True,
                  save_html: bool = True, save_md: bool = True) -> dict[str, str]:
    """Process a single image through the full LlamaCloud OCR pipeline.

    Args:
        api_key: LlamaCloud API key.
        image_path: Path to the image file to process.
        output_dir: Directory where output files should be saved.
        save_word: Whether to save a Word document.
        save_excel: Whether to save an Excel file.
        save_html: Whether to save an HTML file.
        save_md: Whether to save a Markdown file.

    Returns:
        Dictionary with keys 'html', 'md', 'word', 'excel' (if saved) mapping to file paths.

    Raises:
        Exception: If any step in the processing fails.
    """
    # 1. Call LlamaCloud API
    _, base_url = load_api_key()  # We'll get the base_url from config, but we already have api_key
    # Actually, we need to get both from load_api_key, but we already have api_key as param
    # Let's re-call load_api_key to get base_url (it's cheap) or we could have returned both
    # For simplicity, we'll call it again (or we could change the function to accept both)
    # But to avoid breaking the signature, we'll get base_url inside
    _, base_url = load_api_key()
    response_json = call_llamacloud_api(api_key, image_path, base_url)
    
    # 2. Extract markdown content
    markdown_content = extract_markdown_content(response_json)
    
    results = {}
    
    # 3. Save HTML if requested
    if save_html:
        html_path = save_as_html(markdown_content, image_path, output_dir)
        results['html'] = html_path
    
    # 4. Save Markdown if requested
    if save_md:
        md_path = save_as_md(markdown_content, image_path, output_dir)
        results['md'] = md_path
    
    # 5. Extract table and save Word/Excel if requested
    df = None
    if save_word or save_excel:
        # Try to extract table from the markdown content
        df = extract_table_to_dataframe(markdown_content)
        if df is not None and not df.empty:
            # Filter out rows where the first column is empty or all NaN
            if df.shape[1] > 0:
                first_col = df.iloc[:, 0]
                # Remove rows where first column is NaN or empty string
                mask = first_col.notna() & (first_col != '')
                df = df[mask].copy()
                # Also drop rows that are all NaN
                df = df.dropna(how='all')
    
    if save_word and df is not None and not df.empty:
        # Get text before the table for header info
        before_text, _ = split_text_around_table(markdown_content)
        word_path = save_as_word(df, image_path, output_dir, before_text)
        results['word'] = word_path
    
    if save_excel and df is not None and not df.empty:
        excel_path = save_as_excel(df, image_path, output_dir)
        results['excel'] = excel_path
    
    return results