import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Checkbutton, IntVar
import requests
import pandas as pd
from docx import Document
from docx.shared import Cm
import os
import re
import io
import html
import json
from datetime import datetime

class NanoNetsOCRApp:
    def __init__(self, root):
        self.root = root
        self.root.title("NanoNets OCR 表格提取工具")
        self.root.geometry("650x670")
        self.root.configure(bg='black')

        fg_color = 'white'
        bg_color = 'black'
        entry_bg = '#333333'
        btn_bg = '#555555'

        title = tk.Label(root, text="NanoNets OCR 表格提取工具", font=('Arial', 16, 'bold'),
                         fg=fg_color, bg=bg_color)
        title.pack(pady=10)

        # 模型資訊（API Key 來自外部 JSON）
        frame_key = tk.Frame(root, bg=bg_color)
        frame_key.pack(pady=5, fill='x', padx=20)
        tk.Label(frame_key, text="模型：nanonets", fg=fg_color, bg=bg_color,
                 font=('Arial', 10, 'bold')).pack(side='left')
        tk.Label(frame_key, text="（API Key 由 WebOcrAPI.json 提供）", fg='gray', bg=bg_color,
                 font=('Arial', 9)).pack(side='left', padx=5)

        # 輸出目錄選擇
        frame_out = tk.Frame(root, bg=bg_color)
        frame_out.pack(pady=5, fill='x', padx=20)
        tk.Label(frame_out, text="輸出目錄：", fg=fg_color, bg=bg_color,
                 font=('Arial', 10)).pack(side='left')
        self.dir_entry = tk.Entry(frame_out, width=40, bg=entry_bg, fg=fg_color)
        self.dir_entry.pack(side='left', padx=5)
        self.dir_entry.insert(0, r"E:\DiskCUse\HFDownloads")
        tk.Button(frame_out, text="瀏覽", command=self.select_dir,
                  bg=btn_bg, fg=fg_color).pack(side='left', padx=5)

        # 圖片選擇
        frame_file = tk.Frame(root, bg=bg_color)
        frame_file.pack(pady=10, fill='x', padx=20)
        self.file_label = tk.Label(frame_file, text="未選擇圖片", fg='gray', bg=bg_color,
                                   width=40, anchor='w')
        self.file_label.pack(side='left')
        tk.Button(frame_file, text="選擇圖片", command=self.select_file,
                  bg=btn_bg, fg=fg_color, padx=10).pack(side='right')

        # 輸出格式勾選框
        frame_options = tk.LabelFrame(root, text="儲存格式選項", bg=bg_color, fg=fg_color,
                                      font=('Arial', 10))
        frame_options.pack(pady=10, fill='x', padx=20)

        self.var_word = IntVar(value=1)
        self.var_excel = IntVar(value=0)
        self.var_html = IntVar(value=1)
        self.var_md = IntVar(value=1)

        chk_word = Checkbutton(frame_options, text="儲存 WORD 檔", variable=self.var_word,
                               bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_word.pack(side='left', padx=10)

        chk_excel = Checkbutton(frame_options, text="儲存 EXCEL 檔", variable=self.var_excel,
                                bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_excel.pack(side='left', padx=10)

        chk_html = Checkbutton(frame_options, text="儲存原始 HTML 內容", variable=self.var_html,
                               bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_html.pack(side='left', padx=10)

        chk_md = Checkbutton(frame_options, text="儲存 MD 文件", variable=self.var_md,
                             bg=bg_color, fg=fg_color, selectcolor='black', font=('Arial', 10))
        chk_md.pack(side='left', padx=10)

        # 開始按鈕
        self.process_btn = tk.Button(root, text="🚀 開始辨識並轉換", command=self.process,
                                     bg='#0066cc', fg=fg_color, font=('Arial', 12, 'bold'),
                                     padx=20, pady=5)
        self.process_btn.pack(pady=15)

        # 日誌區域
        self.log_area = scrolledtext.ScrolledText(root, height=18, bg='#1a1a1a',
                                                  fg='#00ff00', font=('Consolas', 10))
        self.log_area.pack(fill='both', expand=True, padx=20, pady=10)
        self.log_area.insert(tk.END, "就緒。API Key 將從 WebOcrAPI.json 讀取。請選擇圖片。\n")
        self.log_area.config(state='disabled')

        self.image_path = None

    def select_dir(self):
        path = filedialog.askdirectory(title="選擇輸出目錄")
        if path:
            self.dir_entry.delete(0, tk.END)
            self.dir_entry.insert(0, path)

    def select_file(self):
        path = filedialog.askopenfilename(
            title="選擇圖片檔案",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.bmp *.tiff")]
        )
        if path:
            self.image_path = path
            self.file_label.config(text=os.path.basename(path), fg='white')
            self.log("已選擇檔案：" + path)

    def log(self, msg):
        self.log_area.config(state='normal')
        self.log_area.insert(tk.END, msg + "\n")
        self.log_area.see(tk.END)
        self.log_area.config(state='disabled')

    def load_api_key(self):
        """從 WebOcrAPI.json 讀取 nanonets 模型的 API Key"""
        json_path = os.path.join(os.path.dirname(__file__), "WebOcrAPI.json")
        if not os.path.exists(json_path):
            raise FileNotFoundError(f"找不到設定檔：{json_path}\n請將 WebOcrAPI.json 放在程式相同目錄下。")

        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                configs = json.load(f)
        except json.JSONDecodeError as e:
            raise ValueError(f"WebOcrAPI.json 格式錯誤：{e}")

        if not isinstance(configs, list):
            raise ValueError("WebOcrAPI.json 的內容應為陣列格式。")

        target = None
        for cfg in configs:
            if cfg.get("ModelName") == "nanonets":
                target = cfg
                break

        if target is None:
            raise ValueError("在 WebOcrAPI.json 中找不到 ModelName 為 'nanonets' 的設定。")

        api_key = target.get("api_key", "").strip()
        if not api_key:
            raise ValueError("nanonets 的 API Key 為空，請編輯 WebOcrAPI.json 填入有效的密鑰。")

        return api_key

    def process(self):
        try:
            api_key = self.load_api_key()
        except Exception as e:
            messagebox.showerror("API Key 讀取失敗", str(e))
            return

        if not self.image_path or not os.path.exists(self.image_path):
            messagebox.showerror("錯誤", "請先選擇有效的圖片檔案")
            return

        if not (self.var_word.get() or self.var_excel.get() or self.var_html.get() or self.var_md.get()):
            messagebox.showerror("錯誤", "請至少勾選一種輸出格式")
            return

        out_dir = self.dir_entry.get().strip()
        if not out_dir:
            out_dir = os.getcwd()
        if not os.path.exists(out_dir):
            try:
                os.makedirs(out_dir)
            except Exception as e:
                messagebox.showerror("錯誤", f"無法建立輸出目錄：{e}")
                return

        self.process_btn.config(state='disabled')
        self.log("開始處理...")

        try:
            self.log("正在上傳圖片並呼叫 API...")
            response_json = self.call_nanonets_api(api_key, self.image_path)
            self.log("API 呼叫成功，正在解析 JSON...")

            md_content = self.extract_markdown_content(response_json)
            self.log(f"提取到的內容長度：{len(md_content)} 字元")

            base_name = os.path.splitext(os.path.basename(self.image_path))[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            # 儲存原始 HTML / MD
            if self.var_html.get():
                html_filename = f"{base_name}_{timestamp}.html"
                html_path = os.path.join(out_dir, html_filename)
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                self.log(f"原始 HTML 已儲存：{html_path}")

            if self.var_md.get():
                md_filename = f"{base_name}_{timestamp}.md"
                md_path = os.path.join(out_dir, md_filename)
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                self.log(f"MD 文件已儲存：{md_path}")

            # 解析表格
            df = None
            try:
                self.log("正在解析表格（Markdown）...")
                df = self.extract_table_to_dataframe(md_content)
                self.log(f"解析成功，表格尺寸：{df.shape[0]} 列 x {df.shape[1]} 欄")
            except Exception as e:
                self.log(f"警告：表格解析失敗 - {str(e)}")

            # 提取表格前的文字，用於自動命名
            before_text, _ = self._split_md_around_table(md_content)
            word_filename = self._generate_word_filename(before_text, base_name, timestamp)
            word_path = os.path.join(out_dir, word_filename) if self.var_word.get() else ""

            # 儲存 Excel
            if self.var_excel.get() and df is not None:
                excel_filename = f"{base_name}_{timestamp}.xlsx"
                excel_path = os.path.join(out_dir, excel_filename)
                df.to_excel(excel_path, index=False)
                self.log(f"Excel 已儲存：{excel_path}")
            elif self.var_excel.get() and df is None:
                self.log("略過 Excel：未找到表格資料")

            # 儲存 Word（含自動命名、邊界設定、標頭合併）
            if self.var_word.get() and df is not None:
                self.save_as_word(df, word_path, md_content)
                self.log(f"Word 已儲存：{word_path}")
            elif self.var_word.get() and df is None:
                self.log("略過 Word：未找到表格資料")

            if df is None:
                messagebox.showinfo("完成", "處理完成！\n注意：未從圖片中辨識到表格，僅儲存了原始內容（HTML/MD）。")
            else:
                messagebox.showinfo("完成", f"轉換完成！\n檔案已儲存至：\n{out_dir}")

        except Exception as e:
            error_msg = f"錯誤：{str(e)}"
            self.log(error_msg)
            messagebox.showerror("處理失敗", error_msg)

        finally:
            self.process_btn.config(state='normal')

    def call_nanonets_api(self, api_key, image_path):
        url = "https://extraction-api.nanonets.com/api/v1/extract/sync"
        headers = {"Authorization": f"Bearer {api_key}"}
        files = {"file": (os.path.basename(image_path), open(image_path, 'rb'))}
        data = {"output_format": "markdown"}
        response = requests.post(url, headers=headers, files=files, data=data)
        if response.status_code != 200:
            raise Exception(f"API 回傳錯誤 (HTTP {response.status_code})：{response.text}")
        return response.json()

    def extract_markdown_content(self, json_data):
        try:
            content = json_data.get("result", {}).get("markdown", {}).get("content")
            if content is None:
                content = json_data.get("markdown", {}).get("content")
            if content is None:
                raise ValueError("無法從 JSON 中找到 markdown content")
            return content
        except Exception as e:
            raise Exception(f"解析 JSON 失敗：{e}\n原始回應：{json_data}")

    def _split_md_around_table(self, md_text):
        """從 Markdown 文字中找出第一個表格，回傳 (before_text, after_text)"""
        lines = md_text.splitlines()
        table_start = None
        table_end = None
        in_table = False

        for i, line in enumerate(lines):
            if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
                if not in_table:
                    in_table = True
                    table_start = i
            else:
                if in_table:
                    table_end = i
                    break

        if table_start is None:
            return md_text, ""

        before_lines = lines[:table_start]
        after_lines = lines[table_end:] if table_end is not None else []
        before_text = "\n".join(before_lines).strip()
        after_text = "\n".join(after_lines).strip()
        return before_text, after_text

    def _generate_word_filename(self, before_text, fallback_base, timestamp):
        """從 before_text 嘗試提取日期與庫別以組成檔名（支援 4-2 庫）"""
        date_str = None
        warehouse_str = None

        # 日期：支援中文年月日，中間可有空白
        date_match = re.search(r'(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', before_text)
        if date_match:
            y, m, d = date_match.groups()
            date_str = f"{y}年{m}月{d}日"

        # 庫別：支援純數字或帶連字符（例如 6、4-2、10-1）
        wh_match = re.search(r'\((\d{1,3}(?:-\d{1,2})?)\s*庫\)', before_text)
        if wh_match:
            warehouse_str = f"({wh_match.group(1)}庫)"

        if date_str and warehouse_str:
            return f"{date_str}{warehouse_str}.docx"
        else:
            return f"{fallback_base}_{timestamp}.docx"

    def _merge_title_info(self, lines):
        """合併標頭資訊行（庫別、日期、填單者、單位）"""
        if not lines:
            return lines

        new_lines = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if '庫)' in line and (line.startswith('#') or '(' in line):
                combined_parts = [line]
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if any(next_line.startswith(prefix) for prefix in [
                        '**日期：**', '日期：',
                        '**填單者：**', '填單者：',
                        '**單位：**', '單位：'
                    ]):
                        combined_parts.append(next_line)
                        j += 1
                    else:
                        break
                merged_line = '     '.join(combined_parts)
                new_lines.append(merged_line)
                i = j
            else:
                new_lines.append(line)
                i += 1
        return new_lines

    def extract_table_to_dataframe(self, text):
        md_table = self._extract_markdown_table(text)
        if md_table:
            try:
                df = pd.read_markdown(io.StringIO(md_table))
                return df
            except (ImportError, AttributeError):
                return self._parse_markdown_table(md_table)
            except Exception:
                return self._parse_markdown_table(md_table)
        raise ValueError("未找到任何 Markdown 表格")

    def _extract_markdown_table(self, text):
        lines = text.splitlines()
        table_lines = []
        in_table = False
        for line in lines:
            if '|' in line and (line.strip().startswith('|') or line.strip().endswith('|')):
                if not in_table:
                    in_table = True
                table_lines.append(line)
            else:
                if in_table:
                    if len(table_lines) >= 3:
                        break
                    else:
                        table_lines = []
                        in_table = False
        if len(table_lines) >= 3:
            return '\n'.join(table_lines)
        return None

    def _parse_markdown_table(self, md_text):
        lines = md_text.splitlines()
        data_lines = []
        for line in lines:
            if re.search(r'^\s*\|[\s\-:]+?\|', line):
                continue
            if '|' in line:
                stripped = line.strip('|')
                cells = [cell.strip() for cell in stripped.split('|')]
                data_lines.append(cells)
        if not data_lines:
            raise ValueError("沒有有效的資料行")
        header = data_lines[0]
        data = data_lines[1:]
        max_cols = max(len(row) for row in data_lines)
        for row in data:
            while len(row) < max_cols:
                row.append('')
        while len(header) < max_cols:
            header.append('')
        return pd.DataFrame(data, columns=header[:max_cols])

    def save_as_word(self, df, filename, md_content):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        before_text, after_text = self._split_md_around_table(md_content)
        if before_text:
            before_lines = [line.strip() for line in before_text.split('\n') if line.strip()]
            merged_lines = self._merge_title_info(before_lines)
            for line in merged_lines:
                doc.add_paragraph(line)

        if df.shape[1] > 0:
            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            table.style = 'Table Grid'
            for j, col_name in enumerate(df.columns):
                table.cell(0, j).text = str(col_name)
            for i, row in df.iterrows():
                for j, value in enumerate(row):
                    cell_text = "" if pd.isna(value) else str(value)
                    table.cell(i + 1, j).text = cell_text

        if after_text:
            doc.add_paragraph()
            for line in after_text.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

        doc.save(filename)

    def _clean_html_text(self, raw_text):
        decoded = html.unescape(raw_text)
        plain = re.sub(r'<[^>]+>', '', decoded)
        lines = [line.strip() for line in plain.splitlines() if line.strip()]
        return '\n'.join(lines)

    def extract_text_before_first_table(self, html_str):
        match = re.search(r'<table.*?>', html_str, re.IGNORECASE)
        if not match:
            return ""
        before = html_str[:match.start()]
        return self._clean_html_text(before)

    def extract_text_after_first_table(self, html_str):
        match = re.search(r'<table.*?>.*?</table>', html_str, re.DOTALL | re.IGNORECASE)
        if not match:
            return ""
        after = html_str[match.end():]
        return self._clean_html_text(after)


if __name__ == "__main__":
    root = tk.Tk()
    app = NanoNetsOCRApp(root)
    root.mainloop()